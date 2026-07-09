"""
Herramientas de conversión de archivos y gestión de PDFs.
"""
import os
import json
import logging
import requests

from django.conf import settings
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from ._utils import _convertir_markdown_fallback

logger = logging.getLogger(__name__)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def convertir_markdown(request):
    """
    Convierte archivos (PDF, DOCX, XLSX, PPTX, etc.) a Markdown usando MarkItDown.
    """
    import tempfile

    try:
        # Verificar si hay archivo
        if 'archivo' not in request.FILES:
            return Response({'error': 'No se proporcionó archivo'}, status=400)

        archivo = request.FILES['archivo']

        # Extensiones permitidas
        extensiones_permitidas = ['.pdf', '.docx', '.xlsx', '.pptx', '.doc', '.xls', '.ppt', '.html', '.txt', '.csv', '.json', '.xml']
        ext = os.path.splitext(archivo.name)[1].lower()

        if ext not in extensiones_permitidas:
            return Response({
                'error': f'Extensión no soportada: {ext}',
                'soportadas': extensiones_permitidas
            }, status=400)

        # Guardar archivo temporalmente
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_input:
            for chunk in archivo.chunks():
                tmp_input.write(chunk)
            tmp_input_path = tmp_input.name

        logger.info(f"[MARKITDOWN] Archivo guardado: {tmp_input_path}")

        try:
            markdown = None
            engine = None

            # Intentar usar markitdown como módulo Python.
            try:
                from markitdown import MarkItDown
                md = MarkItDown()
                result = md.convert(tmp_input_path)
                markdown = result.text_content
                engine = 'markitdown'
                logger.info("[MARKITDOWN] Usando módulo Python markitdown")
            except ImportError:
                logger.warning("[MARKITDOWN] Módulo no disponible, usando fallback nativo")
            except Exception as e:
                logger.warning(f"[MARKITDOWN] Conversión con módulo falló: {e}")

            # Fallback: ejecutar CLI.
            if not markdown:
                import subprocess
                import sys
                try:
                    logger.info("[MARKITDOWN] Intentando usar python -m markitdown")
                    resultado = subprocess.run(
                        [sys.executable, '-m', 'markitdown', tmp_input_path],
                        capture_output=True,
                        text=True,
                        timeout=60
                    )
                    if resultado.returncode == 0:
                        markdown = resultado.stdout
                        engine = 'markitdown-cli'
                    else:
                        logger.warning(f"[MARKITDOWN] CLI falló: {resultado.stderr}")
                except Exception as e:
                    logger.warning(f"[MARKITDOWN] CLI no disponible: {e}")

            # Fallback nativo de backend para formatos comunes.
            if not markdown:
                markdown = _convertir_markdown_fallback(tmp_input_path, ext)
                engine = 'fallback'
                logger.info(f"[MARKITDOWN] Conversión fallback usada para {ext}")

            # Limpiar archivo temporal
            try:
                os.unlink(tmp_input_path)
            except Exception:
                pass

            # Extraer metadatos básicos
            metadatos = {
                'nombre_original': archivo.name,
                'extension': ext,
                'tamaño_bytes': archivo.size,
                'lineas_markdown': len(markdown.split('\n')),
                'caracteres': len(markdown),
                'engine': engine,
            }

            logger.info(f"[MARKITDOWN] Conversión exitosa: {archivo.name}")

            return Response({
                'markdown': markdown,
                'metadatos': metadatos,
                'exitoso': True
            })

        except Exception as e:
            try:
                os.unlink(tmp_input_path)
            except Exception:
                pass

            error_msg = str(e)
            logger.error(f"[MARKITDOWN] Error: {error_msg}")

            if 'markitdown' in error_msg.lower() or 'module' in error_msg.lower():
                return Response({
                    'error': 'No se pudo convertir el archivo con MarkItDown ni con fallback',
                    'detalle': error_msg,
                    'instrucciones': 'Instala dependencias y reinicia backend: pip install -r backend/requirements.txt'
                }, status=503)

            if 'Fallback no soporta extensión' in error_msg:
                return Response({
                    'error': 'Tipo de archivo no soportado sin MarkItDown',
                    'detalle': error_msg,
                    'instrucciones': 'Instala MarkItDown para convertir este tipo de archivo'
                }, status=422)

            return Response({
                'error': 'Error al convertir archivo',
                'detalle': error_msg
            }, status=500)

    except Exception as e:
        logger.error(f"[MARKITDOWN] Error general: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return Response({
            'error': 'Error interno del servidor',
            'detalle': str(e)
        }, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def convertir_archivo(request):
    """
    Convierte archivos entre diferentes formatos.

    Espera:
    - archivo: file upload
    - formato_destino: extensión de destino (pdf, docx, xlsx, csv, txt, html)

    Retorna:
    - archivo_base64: contenido del archivo convertido
    - nombre_archivo: nombre sugerido para descargar
    - mime_type: tipo MIME del archivo
    """
    import tempfile
    import base64
    import subprocess
    import sys

    try:
        if 'archivo' not in request.FILES:
            return Response({'error': 'No se proporcionó archivo'}, status=400)

        if 'formato_destino' not in request.POST:
            return Response({'error': 'No se especificó formato de destino'}, status=400)

        archivo = request.FILES['archivo']
        formato_destino = request.POST['formato_destino'].lower().strip('.')

        # Detectar formato origen
        ext_origen = os.path.splitext(archivo.name)[1].lower()
        formato_origen = ext_origen.lstrip('.')
        nombre_base = os.path.splitext(archivo.name)[0]

        logger.info(f"[CONVERTIR] {formato_origen} -> {formato_destino}: {archivo.name}")

        # Guardar archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext_origen) as tmp_input:
            for chunk in archivo.chunks():
                tmp_input.write(chunk)
            input_path = tmp_input.name

        output_path = input_path.replace(ext_origen, f'.{formato_destino}')

        try:
            resultado = None

            # === PDF a DOCX ===
            if formato_origen == 'pdf' and formato_destino == 'docx':
                try:
                    from pdf2docx import Converter
                    cv = Converter(input_path)
                    cv.convert(output_path, start=0, end=None)
                    cv.close()
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'pdf2docx no está instalado',
                        'instrucciones': 'pip install pdf2docx'
                    }, status=503)
                except BaseException as e:
                    return Response({
                        'error': 'No se pudo convertir el PDF. El archivo puede estar dañado o tener un formato no compatible.',
                        'detalle': str(e)
                    }, status=422)
            # === PDF a TXT (usando markitdown) ===
            elif formato_origen == 'pdf' and formato_destino == 'txt':
                try:
                    from markitdown import MarkItDown
                    md = MarkItDown()
                    result = md.convert(input_path)
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(result.text_content)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'markitdown no está instalado',
                        'instrucciones': 'pip install markitdown'
                    }, status=503)

            # === DOCX a PDF (usando reportlab - pura Python, sin dependencias externas) ===
            elif formato_origen == 'docx' and formato_destino == 'pdf':
                try:
                    from docx import Document
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

                    # Leer documento Word
                    doc = Document(input_path)

                    # Crear PDF
                    pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []

                    # Convertir cada párrafo
                    for para in doc.paragraphs:
                        if para.text.strip():
                            # Detectar si es título (estilo Heading)
                            style_name = para.style.name if para.style else 'Normal'
                            if 'Heading 1' in style_name:
                                style = ParagraphStyle(
                                    'Heading1',
                                    parent=styles['Heading1'],
                                    fontSize=18,
                                    spaceAfter=12
                                )
                            elif 'Heading 2' in style_name:
                                style = ParagraphStyle(
                                    'Heading2',
                                    parent=styles['Heading2'],
                                    fontSize=14,
                                    spaceAfter=10
                                )
                            else:
                                style = styles['Normal']

                            story.append(Paragraph(para.text, style))
                            story.append(Spacer(1, 6))

                    pdf_doc.build(story)
                    resultado = output_path

                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install python-docx reportlab'
                    }, status=503)

            # === DOCX a HTML ===
            elif formato_origen == 'docx' and formato_destino == 'html':
                try:
                    from mammoth import convert_to_html
                    with open(input_path, 'rb') as docx_file:
                        result = convert_to_html(docx_file)
                        with open(output_path, 'w', encoding='utf-8') as f:
                            f.write(result.value)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'mammoth no está instalado',
                        'instrucciones': 'pip install mammoth'
                    }, status=503)

            # === DOCX a TXT ===
            elif formato_origen == 'docx' and formato_destino == 'txt':
                try:
                    from docx import Document
                    doc = Document(input_path)
                    texto = '\n'.join([para.text for para in doc.paragraphs])
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(texto)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'python-docx no está instalado',
                        'instrucciones': 'pip install python-docx'
                    }, status=503)

            # === XLSX a CSV ===
            elif formato_origen == 'xlsx' and formato_destino == 'csv':
                import csv
                import openpyxl
                wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
                ws = wb.active
                with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    for row in ws.iter_rows(values_only=True):
                        writer.writerow(['' if v is None else v for v in row])
                wb.close()
                resultado = output_path

            # === XLSX a PDF (usando reportlab - pura Python) ===
            elif formato_origen == 'xlsx' and formato_destino == 'pdf':
                try:
                    import openpyxl
                    from reportlab.lib.pagesizes import letter, landscape
                    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
                    from reportlab.lib import colors

                    wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
                    ws = wb.active
                    rows = list(ws.iter_rows(values_only=True))
                    wb.close()

                    # Convertir celdas a str para reportlab (no acepta None ni tipos numéricos en Table)
                    data = [['' if v is None else str(v) for v in row] for row in rows]

                    pdf_doc = SimpleDocTemplate(output_path, pagesize=landscape(letter))
                    table = Table(data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#001e33')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                    ]))
                    pdf_doc.build([table])
                    resultado = output_path

                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install openpyxl reportlab'
                    }, status=503)

            # === XLSX a TXT ===
            elif formato_origen == 'xlsx' and formato_destino == 'txt':
                import openpyxl
                wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
                ws = wb.active
                rows = list(ws.iter_rows(values_only=True))
                wb.close()
                if rows:
                    # Calcular anchos de columna para alinear el texto
                    n_cols = max(len(r) for r in rows)
                    col_widths = [
                        max(len(str(rows[r][c]) if c < len(rows[r]) and rows[r][c] is not None else '') for r in range(len(rows)))
                        for c in range(n_cols)
                    ]
                    with open(output_path, 'w', encoding='utf-8') as f:
                        for row in rows:
                            line = '  '.join(
                                str(v if v is not None else '').ljust(col_widths[i])
                                for i, v in enumerate(row)
                            )
                            f.write(line.rstrip() + '\n')
                resultado = output_path

            # === CSV a XLSX ===
            elif formato_origen == 'csv' and formato_destino == 'xlsx':
                import csv
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                with open(input_path, 'r', encoding='utf-8-sig', newline='') as f:
                    for row in csv.reader(f):
                        ws.append(row)
                wb.save(output_path)
                resultado = output_path

            # === TXT a DOCX ===
            elif formato_origen == 'txt' and formato_destino == 'docx':
                try:
                    from docx import Document
                    doc = Document()
                    with open(input_path, 'r', encoding='utf-8') as f:
                        for linea in f:
                            doc.add_paragraph(linea.strip())
                    doc.save(output_path)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'python-docx no está instalado',
                        'instrucciones': 'pip install python-docx'
                    }, status=503)

            # === TXT a PDF ===
            elif formato_origen == 'txt' and formato_destino == 'pdf':
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet
                    doc = SimpleDocTemplate(output_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                    with open(input_path, 'r', encoding='utf-8') as f:
                        for linea in f:
                            txt = linea.rstrip()
                            if txt:
                                story.append(Paragraph(txt, styles['Normal']))
                                story.append(Spacer(1, 4))
                            else:
                                story.append(Spacer(1, 12))
                    doc.build(story)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'reportlab no está instalado',
                        'instrucciones': 'pip install reportlab'
                    }, status=503)

            # === HTML a PDF (usando reportlab con html.parser) ===
            elif formato_origen == 'html' and formato_destino == 'pdf':
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet
                    from html.parser import HTMLParser

                    # Parser simple para extraer texto del HTML
                    class HTMLTextExtractor(HTMLParser):
                        def __init__(self):
                            super().__init__()
                            self.texts = []
                            self.current_tag = None

                        def handle_starttag(self, tag, attrs):
                            self.current_tag = tag

                        def handle_endtag(self, tag):
                            if tag in ('p', 'div', 'h1', 'h2', 'h3', 'br'):
                                self.texts.append('\n')
                            self.current_tag = None

                        def handle_data(self, data):
                            if data.strip():
                                self.texts.append(data.strip())

                    # Leer HTML
                    with open(input_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()

                    parser = HTMLTextExtractor()
                    parser.feed(html_content)
                    text = ' '.join(parser.texts)

                    # Crear PDF
                    pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []

                    for line in text.split('\n'):
                        if line.strip():
                            story.append(Paragraph(line.strip(), styles['Normal']))
                            story.append(Spacer(1, 6))

                    pdf_doc.build(story)
                    resultado = output_path

                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install reportlab'
                    }, status=503)

            # === HTML a DOCX ===
            elif formato_origen == 'html' and formato_destino == 'docx':
                try:
                    from htmldocx import HtmlToDocx
                    parser = HtmlToDocx()
                    parser.parse_html_file(input_path, output_path)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'htmldocx no está instalado',
                        'instrucciones': 'pip install htmldocx'
                    }, status=503)

            # === IMÁGENES a PDF (png, jpg, jpeg, gif, bmp, webp) ===
            elif formato_origen in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp') and formato_destino == 'pdf':
                try:
                    from PIL import Image
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Image as RLImage, Spacer
                    from reportlab.lib.units import inch

                    # Abrir imagen y obtener dimensiones
                    img = Image.open(input_path)
                    img_width, img_height = img.size

                    # Calcular tamaño para ajustar a página
                    page_width, page_height = letter
                    max_width = page_width - 2 * inch
                    max_height = page_height - 2 * inch

                    # Escalar proporcionalmente
                    scale = min(max_width / img_width, max_height / img_height, 1.0)
                    final_width = img_width * scale
                    final_height = img_height * scale

                    # Crear PDF
                    pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
                    story = []
                    story.append(Spacer(1, 0.5 * inch))
                    story.append(RLImage(input_path, width=final_width, height=final_height))
                    pdf_doc.build(story)
                    resultado = output_path

                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install Pillow reportlab'
                    }, status=503)

            else:
                return Response({
                    'error': f'Conversión no soportada: {formato_origen} -> {formato_destino}'
                }, status=400)

            # Leer archivo convertido y codificar en base64
            if resultado and os.path.exists(resultado):
                with open(resultado, 'rb') as f:
                    contenido = f.read()
                    archivo_base64 = base64.b64encode(contenido).decode('utf-8')

                # Determinar MIME type
                mime_types = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    'csv': 'text/csv',
                    'txt': 'text/plain',
                    'html': 'text/html',
                }

                # Limpiar archivos temporales
                try:
                    os.unlink(input_path)
                    os.unlink(resultado)
                except Exception:
                    pass

                logger.info(f"[CONVERTIR] Éxito: {archivo.name} -> {nombre_base}-convertido.{formato_destino}")

                return Response({
                    'archivo_base64': archivo_base64,
                    'nombre_archivo': f'{nombre_base}-convertido.{formato_destino}',
                    'mime_type': mime_types.get(formato_destino, 'application/octet-stream'),
                    'exitoso': True
                })
            else:
                return Response({
                    'error': 'Error al generar archivo convertido'
                }, status=500)

        except Exception as e:
            # Limpiar archivos temporales
            try:
                os.unlink(input_path)
                if os.path.exists(output_path):
                    os.unlink(output_path)
            except Exception:
                pass

            logger.error(f"[CONVERTIR] Error: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())

            return Response({
                'error': 'Error al convertir archivo',
                'detalle': str(e)
            }, status=500)

    except Exception as e:
        logger.error(f"[CONVERTIR] Error general: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return Response({
            'error': 'Error interno del servidor',
            'detalle': str(e)
        }, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def gestor_pdf(request):
    """
    Endpoint para gestionar operaciones con PDFs:
    - fusionar: Unir múltiples PDFs en uno solo
    - dividir: Separar páginas de un PDF
    - rotar: Rotar páginas
    - extraer: Extraer páginas específicas
    - marca: Agregar marca de agua de texto
    - proteger: Agregar contraseña
    - desbloquear: Quitar contraseña
    """
    try:
        import tempfile
        import base64
        import io

        payload = request.data if isinstance(request.data, dict) else {}
        herramienta = (payload.get('herramienta') or request.POST.get('herramienta') or '').strip().lower()
        if not herramienta:
            return Response({'error': 'No se especificó la herramienta'}, status=400)

        resultado_archivos = []

        try:
            cantidad = int(request.POST.get('cantidad_archivos', payload.get('cantidad_archivos', 0)) or 0)
        except (TypeError, ValueError):
            return Response({'error': 'cantidad_archivos inválido'}, status=400)

        if cantidad <= 0:
            return Response({'error': 'No se proporcionaron archivos'}, status=400)

        archivos_pdf = [request.FILES[f'archivo_{i}'] for i in range(cantidad) if f'archivo_{i}' in request.FILES]
        if not archivos_pdf:
            archivos_pdf = list(request.FILES.values())
        if not archivos_pdf:
            return Response({'error': 'No se encontraron archivos PDF'}, status=400)

        logger.info(f"[GESTOR PDF] Herramienta: {herramienta}, Archivos: {len(archivos_pdf)}")

        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            return Response({'error': 'pypdf no está instalado', 'instrucciones': 'pip install pypdf'}, status=503)

        # FUSIONAR PDFs
        if herramienta == 'fusionar':
            writer = PdfWriter()
            for archivo in archivos_pdf:
                reader = PdfReader(archivo)
                for page in reader.pages:
                    writer.add_page(page)
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            with open(output_path, 'rb') as f:
                contenido = f.read()
            resultado_archivos.append({'nombre': 'fusionado.pdf', 'contenido_base64': base64.b64encode(contenido).decode('utf-8')})
            os.unlink(output_path)

        # DIVIDIR PDF
        elif herramienta == 'dividir':
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            for i in range(len(reader.pages)):
                writer = PdfWriter()
                writer.add_page(reader.pages[i])
                output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
                with open(output_path, 'wb') as f:
                    writer.write(f)
                with open(output_path, 'rb') as f:
                    contenido = f.read()
                resultado_archivos.append({'nombre': f'pagina_{i+1}.pdf', 'contenido_base64': base64.b64encode(contenido).decode('utf-8')})
                os.unlink(output_path)

        # ROTAR PÁGINAS
        elif herramienta == 'rotar':
            try:
                rotacion = int(request.POST.get('rotacion', 90))
            except (TypeError, ValueError):
                return Response({'error': 'Rotación inválida'}, status=400)
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            for page in reader.pages:
                try:
                    page.rotate(rotacion)
                except Exception:
                    page.rotate_clockwise(rotacion)
                writer.add_page(page)
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            with open(output_path, 'rb') as f:
                contenido = f.read()
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({'nombre': f'{nombre_base}_rotado_{rotacion}.pdf', 'contenido_base64': base64.b64encode(contenido).decode('utf-8')})
            os.unlink(output_path)

        # EXTRAER PÁGINAS
        elif herramienta == 'extraer':
            paginas_str = request.POST.get('paginas', '')
            if not paginas_str:
                return Response({'error': 'Debes indicar las páginas a extraer'}, status=400)
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            paginas_a_extraer = set()
            for parte in paginas_str.split(','):
                parte = parte.strip()
                if '-' in parte:
                    try:
                        inicio, fin = map(int, parte.split('-'))
                        paginas_a_extraer.update(range(inicio, fin + 1))
                    except (TypeError, ValueError):
                        continue
                elif parte.isdigit():
                    paginas_a_extraer.add(int(parte))
            for num_pagina in sorted(paginas_a_extraer):
                if 1 <= num_pagina <= len(reader.pages):
                    writer.add_page(reader.pages[num_pagina - 1])
            if len(writer.pages) == 0:
                return Response({'error': 'No se encontraron páginas válidas para extraer'}, status=400)
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            with open(output_path, 'rb') as f:
                contenido = f.read()
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({'nombre': f'{nombre_base}_extraido.pdf', 'contenido_base64': base64.b64encode(contenido).decode('utf-8')})
            os.unlink(output_path)

        # MARCA DE AGUA
        elif herramienta == 'marca':
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas
                from reportlab.lib.utils import ImageReader
            except ImportError:
                return Response({'error': 'reportlab no está instalado', 'instrucciones': 'pip install reportlab'}, status=503)
            texto_marca = request.POST.get('texto', '')
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            for page in reader.pages:
                packet = io.BytesIO()
                c = canvas.Canvas(packet, pagesize=letter)
                if texto_marca:
                    c.saveState()
                    c.setFont('Helvetica', 40)
                    try:
                        c.setFillColorRGB(0.7, 0.7, 0.7, alpha=0.3)
                    except TypeError:
                        c.setFillColorRGB(0.7, 0.7, 0.7)
                    c.translate(300, 400)
                    c.rotate(45)
                    c.drawCentredString(0, 0, texto_marca)
                    c.restoreState()
                c.save()
                packet.seek(0)
                marca_reader = PdfReader(packet)
                page.merge_page(marca_reader.pages[0])
                writer.add_page(page)
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            with open(output_path, 'rb') as f:
                contenido = f.read()
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({'nombre': f'{nombre_base}_con_marca.pdf', 'contenido_base64': base64.b64encode(contenido).decode('utf-8')})
            os.unlink(output_path)

        # PROTEGER PDF
        elif herramienta == 'proteger':
            password = request.POST.get('password', '')
            if not password:
                return Response({'error': 'Se requiere una contraseña'}, status=400)
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            writer.encrypt(password)
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            with open(output_path, 'rb') as f:
                contenido = f.read()
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({'nombre': f'{nombre_base}_protegido.pdf', 'contenido_base64': base64.b64encode(contenido).decode('utf-8')})
            os.unlink(output_path)

        # DESBLOQUEAR PDF
        elif herramienta == 'desbloquear':
            password = request.POST.get('password', '')
            archivo = archivos_pdf[0]
            try:
                reader = PdfReader(archivo)
                if reader.is_encrypted:
                    reader.decrypt(password)
                writer = PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
                with open(output_path, 'wb') as f:
                    writer.write(f)
                with open(output_path, 'rb') as f:
                    contenido = f.read()
                nombre_base = os.path.splitext(archivo.name)[0]
                resultado_archivos.append({'nombre': f'{nombre_base}_desbloqueado.pdf', 'contenido_base64': base64.b64encode(contenido).decode('utf-8')})
                os.unlink(output_path)
            except Exception:
                return Response({'error': 'No se pudo desbloquear el PDF. Verifica la contraseña.'}, status=400)

        else:
            return Response({'error': f'Herramienta no reconocida: {herramienta}'}, status=400)

        logger.info(f"[GESTOR PDF] Éxito: {herramienta}, {len(resultado_archivos)} archivos generados")

        return Response({
            'archivos': resultado_archivos,
            'exitoso': True
        })

    except Exception as e:
        logger.error(f"[GESTOR PDF] Error: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return Response({
            'error': 'Error al procesar PDF',
            'detalle': str(e)
        }, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def descargar_archivo_intranet(request):
    """
    Proxy autenticado para descargar archivos de SharePoint vía n8n.
    Body: { "tipo": "contratos"|"reglamento"|"clientes"|"cursos"|"datos_academicos",
            "archivo": "nombre_del_archivo.pdf" }
    """
    tipo = request.data.get('tipo', '').strip()
    archivo = request.data.get('archivo', '').strip()

    TIPOS_VALIDOS = {'contratos', 'reglamento', 'clientes', 'cursos', 'datos_academicos'}
    if tipo not in TIPOS_VALIDOS:
        return Response({'error': f'tipo inválido. Valores permitidos: {", ".join(TIPOS_VALIDOS)}'}, status=400)
    if not archivo:
        return Response({'error': 'archivo es requerido'}, status=400)

    try:
        from ..n8n_gateway import descargar_intranet
        resp = descargar_intranet(tipo, archivo)
    except RuntimeError as e:
        logger.error(f"[DESCARGAR INTRANET] Config error: {e}")
        return Response({'error': str(e)}, status=503)
    except requests.exceptions.HTTPError as e:
        logger.error(f"[DESCARGAR INTRANET] HTTP error {e.response.status_code}: {e.response.text[:200]}")
        return Response({'error': f'SharePoint devolvió error {e.response.status_code}'}, status=502)
    except requests.exceptions.ConnectionError:
        logger.error('[DESCARGAR INTRANET] No se pudo conectar con n8n')
        return Response({'error': 'No se pudo conectar con el servicio de archivos'}, status=503)

    from django.http import HttpResponse
    ext = archivo.rsplit('.', 1)[-1].lower() if '.' in archivo else ''
    mime_map = {'pdf': 'application/pdf', 'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'}
    content_type = mime_map.get(ext, resp.headers.get('Content-Type', 'application/octet-stream'))
    response = HttpResponse(resp.content, content_type=content_type)
    # inline para que el iframe lo muestre en lugar de descargarlo
    response['Content-Disposition'] = f'inline; filename="{archivo}"'
    return response
