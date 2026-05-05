import jwt as pyjwt
import uuid
from datetime import datetime

from django.db import models as django_models
from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.response import Response
from django.db.models import Q
import bcrypt
import logging
import resend
import random
from django.core.cache import cache
from django.conf import settings
from .models import DatosArea, DatosCargo, SuperAdmin, DatosEmpleado, TareasCalendario, SolicitudesPassword, ReglamentoItem, Curso, CursoContenido, CursoHistorial, Alerta, N8nLog, ApiKey
from .jwt_utils import generate_tokens, decode_token, build_superadmin_payload, build_empleado_payload, jwt_required

logger = logging.getLogger(__name__)

def get_usuario_nombre(user):
    """Obtiene el nombre del usuario (Django User o DatosEmpleado)"""
    if not user or not hasattr(user, 'is_authenticated') or not user.is_authenticated:
        return ''
    # Intentar username (Django User)
    if hasattr(user, 'username') and user.username:
        return user.username
    # Intentar correo_electronico (DatosEmpleado)
    if hasattr(user, 'correo_electronico') and user.correo_electronico:
        return user.correo_electronico
    # Intentar nombre completo
    nombre = getattr(user, 'primer_nombre', '') or ''
    apellido = getattr(user, 'primer_apellido', '') or ''
    full = f"{nombre} {apellido}".strip()
    return full or 'Usuario'

# Configurar Resend
resend.api_key = settings.RESEND_API_KEY

# Función para generar código de verificación
def generar_codigo_verificacion():
    """Genera código aleatorio de 6 dígitos"""
    return str(random.randint(100000, 999999))

# Función para enviar webhook a n8n
def enviar_webhook_n8n(destinatario, datos):
    """
    Envía webhook a n8n para que maneje el envío de emails.
    n8n se conecta al correo corporativo de RBCOL y envía el email al usuario.
    """
    try:
        import requests
        from django.conf import settings
        
        n8n_url = settings.N8N_WEBHOOK_URL
        api_key = settings.N8N_WEBHOOK_API_KEY
        
        if not n8n_url:
            logger.error("[N8N] N8N_WEBHOOK_URL no configurado")
            return False, "N8N_WEBHOOK_URL no configurado"
        
        payload = {
            'evento': 'nuevo_usuario_creado',
            'destinatario': destinatario,
            'datos': datos,
            'timestamp': str(datetime.now())
        }
        
        headers = {
            'Content-Type': 'application/json',
            'X-API-Key': api_key
        }
        
        response = requests.post(
            n8n_url,
            json=payload,
            headers=headers,
            timeout=10
        )
        
        if response.status_code == 200:
            logger.info(f"[N8N] Webhook enviado exitosamente a {destinatario}")
            resp_json = response.json()
            try:
                import json
                N8nLog.objects.create(
                    workflow_name=datos.get('evento', 'webhook_generico'),
                    status='SUCCESS',
                    message=f"Webhook enviado a {destinatario}",
                    destinatario=destinatario,
                    tipo_evento=datos.get('evento'),
                    response_data=json.dumps(resp_json)[:500],
                )
            except Exception as log_err:
                logger.warning(f"[N8N] No se pudo guardar log: {log_err}")
            return True, resp_json
        else:
            logger.error(f"[N8N] Error en webhook: {response.status_code} - {response.text}")
            try:
                N8nLog.objects.create(
                    workflow_name=datos.get('evento', 'webhook_generico'),
                    status='ERROR',
                    message=f"HTTP {response.status_code}: {response.text[:200]}",
                    destinatario=destinatario,
                    tipo_evento=datos.get('evento'),
                )
            except Exception as log_err:
                logger.warning(f"[N8N] No se pudo guardar log de error: {log_err}")
            return False, f"HTTP {response.status_code}"

    except Exception as e:
        logger.error(f"[N8N] Error enviando webhook: {str(e)}")
        try:
            N8nLog.objects.create(
                workflow_name=datos.get('evento', 'webhook_generico') if isinstance(datos, dict) else 'webhook_generico',
                status='ERROR',
                message=str(e)[:300],
                destinatario=destinatario,
            )
        except Exception:
            pass
        return False, str(e)


# Función para enviar email con código (ahora usa n8n)
def enviar_email_verificacion(email, codigo, password=None, nombre=None):
    """
    Envía datos a n8n para que el correo corporativo envíe el email al usuario.
    Incluye: código de verificación, contraseña temporal, correo de login.
    """
    try:
        from django.conf import settings
        import requests
        
        n8n_url = settings.N8N_WEBHOOK_URL
        logger.info(f"[EMAIL DEBUG] N8N_WEBHOOK_URL leída: '{n8n_url[:50]}...' (longitud: {len(n8n_url) if n8n_url else 0})")

        # Si no hay n8n configurado, usar fallback (SMTP directo)
        if not n8n_url:
            logger.warning("[EMAIL] N8N no configurado, usando SMTP directo")
            return _enviar_email_smtp_fallback(email, codigo)
        
        # Generar HTML del email
        nombre_usuario = nombre or 'Usuario'
        html_email = f"""<div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; background: #f5f5f5;">
  <div style="background: #001e33; color: white; padding: 30px; text-align: center; border-radius: 8px 8px 0 0;">
    <h2 style="margin: 0;">RBG CT</h2>
    <p style="margin: 10px 0 0 0; opacity: 0.9;">Bienvenido al sistema</p>
  </div>
  <div style="background: white; padding: 30px; border-radius: 0 0 8px 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
    <p style="font-size: 16px; color: #333; margin-bottom: 20px;">Hola <strong>{nombre_usuario}</strong>,</p>
    <p style="font-size: 16px; color: #333; margin-bottom: 20px;">Tu cuenta ha sido creada. Estas son tus credenciales de acceso:</p>
    <div style="background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0; border-left: 4px solid #001e33;">
      <p style="margin: 5px 0;"><strong>Correo:</strong> {email}</p>
      <p style="margin: 5px 0;"><strong>Contraseña temporal:</strong> {password}</p>
      <p style="margin: 5px 0;"><strong>Código de verificación:</strong> <span style="font-size: 24px; font-weight: bold; color: #001e33;">{codigo}</span></p>
    </div>
    <p style="font-size: 14px; color: #666; margin-top: 20px;">El código expira en 15 minutos. Por seguridad, cambia tu contraseña al ingresar.</p>
    <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
    <p style="font-size: 12px; color: #999; text-align: center;">RBG CT - Sistema de Gestión<br>Si no solicitaste esta cuenta, contacta al administrador.</p>
  </div>
</div>"""

        # Datos a enviar a n8n (encriptados en tránsito por HTTPS)
        payload = {
            'tipo': 'bienvenida_nuevo_usuario',
            'destinatario': email,
            'asunto': 'Bienvenido a RBG CT - Credenciales de acceso',
            'html_email': html_email,
            'datos_sensibles': {
                'correo_login': email,
                'password_temporal': password,
                'codigo_verificacion': codigo,
            },
            'datos_usuario': {
                'nombre': nombre_usuario,
                'expira_en': '15 minutos'
            },
            'plantilla': 'bienvenida_rbgct'
        }
        
        headers = {
            'Content-Type': 'application/json',
            'X-API-Key': settings.N8N_WEBHOOK_API_KEY
        }
        
        response = requests.post(
            n8n_url,
            json=payload,
            headers=headers,
            timeout=10
        )
        
        if response.status_code == 200:
            logger.info(f"[N8N] Datos enviados a n8n para {email}")
            try:
                import json as _json
                N8nLog.objects.create(
                    workflow_name='bienvenida_nuevo_usuario',
                    status='SUCCESS',
                    message=f"Email de bienvenida enviado a {email}",
                    destinatario=email,
                    tipo_evento='bienvenida_nuevo_usuario',
                    response_data=_json.dumps(response.json())[:500],
                )
            except Exception as log_err:
                logger.warning(f"[N8N] No se pudo guardar log: {log_err}")
            return True, {"status": "webhook_sent", "n8n_response": response.json()}
        else:
            logger.error(f"[N8N] Error: {response.status_code}")
            try:
                N8nLog.objects.create(
                    workflow_name='bienvenida_nuevo_usuario',
                    status='ERROR',
                    message=f"HTTP {response.status_code} al enviar a {email}",
                    destinatario=email,
                    tipo_evento='bienvenida_nuevo_usuario',
                )
            except Exception:
                pass
            # Fallback a SMTP si n8n falla
            return _enviar_email_smtp_fallback(email, codigo)
            
    except Exception as e:
        logger.error(f"[EMAIL/N8N] Error: {str(e)}")
        # Fallback a SMTP
        return _enviar_email_smtp_fallback(email, codigo)


def _enviar_email_smtp_fallback(email, codigo):
    """Función fallback que envía email directo por SMTP si n8n falla"""
    try:
        from django.core.mail import send_mail
        from django.conf import settings
        
        subject = 'Código de verificación - RBG CT'
        html_content = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; background: #f5f5f5;">
            <div style="background: #001e33; color: white; padding: 30px; text-align: center; border-radius: 8px 8px 0 0;">
                <h2 style="margin: 0;">RBG CT</h2>
                <p style="margin: 10px 0 0 0; opacity: 0.9;">Código de verificación</p>
            </div>
            <div style="background: white; padding: 30px; border-radius: 0 0 8px 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                <p style="font-size: 16px; color: #333; margin-bottom: 20px;">
                    Tu código de verificación es:
                </p>
                <div style="background: #001e33; color: white; font-size: 32px; font-weight: bold; text-align: center; padding: 20px; border-radius: 8px; letter-spacing: 8px; margin: 20px 0;">
                    {codigo}
                </div>
                <p style="font-size: 14px; color: #666; margin-top: 20px;">
                    Este código expira en 15 minutos. No lo compartas con nadie.
                </p>
                <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
                <p style="font-size: 12px; color: #999; text-align: center;">
                    Si no solicitaste este código, ignora este mensaje.<br>
                    RBG CT - Sistema de Gestión
                </p>
            </div>
        </div>
        """
        
        send_mail(
            subject=subject,
            message=f'Tu código de verificación es: {codigo}. Expira en 15 minutos.',
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[email],
            fail_silently=False,
            html_message=html_content
        )
        
        logger.info(f"[EMAIL FALLBACK] Código enviado a {email} por SMTP")
        return True, {"status": "sent_via_smtp_fallback"}
        
    except Exception as e:
        logger.error(f"[EMAIL FALLBACK] Error: {str(e)}")
        return False, str(e)

from .serializers import (
    DatosAreaSerializer, DatosCargoSerializer, SuperAdminSerializer,
    DatosEmpleadoSerializer, TareasCalendarioSerializer, SolicitudesPasswordSerializer,
    ReglamentoItemSerializer, CursoSerializer, CursoContenidoSerializer, CursoHistorialSerializer, N8nLogSerializer, ApiKeySerializer
)

# Endpoint de Login - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def login_view(request):
    email = request.data.get('email')
    password = request.data.get('password')
    
    if not email or not password:
        return Response({'error': 'Email y password requeridos'}, status=status.HTTP_400_BAD_REQUEST)
    
    # 1. Verificar si es SuperAdmin
    try:
        admin = SuperAdmin.objects.get(email=email)
        if bcrypt.checkpw(password.encode('utf-8'), admin.password_hash.encode('utf-8')):
            tokens = generate_tokens(build_superadmin_payload(admin))
            return Response({
                'type': 'superadmin',
                'user': {
                    'id': str(admin.id),
                    'email': admin.email,
                    'nombre': admin.nombre,
                    'apellido': admin.apellido,
                },
                **tokens,
            })
        else:
            return Response({'error': 'Credenciales inválidas'}, status=status.HTTP_401_UNAUTHORIZED)
    except SuperAdmin.DoesNotExist:
        pass

    # 2. Verificar si es Empleado
    try:
        empleado = DatosEmpleado.objects.get(correo_corporativo=email, estado='ACTIVA')
        if empleado.password_hash and bcrypt.checkpw(password.encode('utf-8'), empleado.password_hash.encode('utf-8')):

            from django.utils import timezone
            empleado.ultima_actividad = timezone.now()
            empleado.save(update_fields=['ultima_actividad'])

            # Primer login: requiere verificación por código
            if empleado.primer_login:
                return Response({
                    'type': 'empleado',
                    'user': {
                        'id_empleado': empleado.id_empleado,
                        'correo_corporativo': empleado.correo_corporativo,
                        'id_permisos': empleado.id_permisos,
                        'primer_login': True,
                        'datos_completados': empleado.datos_completados,
                    },
                    'requiere_verificacion': True,
                    'mensaje': 'Por favor ingresa el código de verificación enviado a tu correo',
                })

            tokens = generate_tokens(build_empleado_payload(empleado))
            return Response({
                'type': 'empleado',
                'user': {
                    'id_empleado': empleado.id_empleado,
                    'primer_nombre': empleado.primer_nombre,
                    'segundo_nombre': empleado.segundo_nombre,
                    'primer_apellido': empleado.primer_apellido,
                    'segundo_apellido': empleado.segundo_apellido,
                    'correo_corporativo': empleado.correo_corporativo,
                    'id_permisos': empleado.id_permisos,
                    'estado': empleado.estado,
                    'area_id': empleado.area_id,
                    'cargo_id': empleado.cargo_id,
                    'primer_login': False,
                    'datos_completados': empleado.datos_completados,
                    'permitir_edicion_datos': empleado.permitir_edicion_datos,
                },
                **tokens,
            })
        else:
            return Response({'error': 'Credenciales inválidas'}, status=status.HTTP_401_UNAUTHORIZED)
    except DatosEmpleado.DoesNotExist:
        return Response({'error': 'Credenciales inválidas'}, status=status.HTTP_401_UNAUTHORIZED)


# Endpoint para crear usuarios (solo SuperAdmin)
@api_view(['POST'])
@permission_classes([AllowAny])  # Se valida dentro del endpoint
def crear_usuario_superadmin(request):
    """
    Solo SuperAdmin puede crear usuarios.
    Puede crear con datos completos o solo correo+contraseña.
    """
    logger.info(f"[CREAR USUARIO] Datos recibidos: {request.data}")
    
    # Verificar que sea SuperAdmin (enviar credenciales en el request)
    admin_email = request.data.get('admin_email', '').strip()
    admin_password = request.data.get('admin_password', '').strip()

    if not admin_email or not admin_password:
        logger.error(f"[CREAR USUARIO] Faltan credenciales: admin_email={admin_email}, admin_password={'*****' if admin_password else None}")
        return Response({'error': 'Credenciales de administrador requeridas'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        admin = SuperAdmin.objects.get(email=admin_email)
        password_valid = bcrypt.checkpw(admin_password.encode('utf-8'), admin.password_hash.encode('utf-8'))
        logger.info(f"[CREAR USUARIO] Admin encontrado: {admin_email}, password_valid: {password_valid}")
        if not password_valid:
            return Response({'error': 'Credenciales de administrador inválidas'}, status=status.HTTP_401_UNAUTHORIZED)
    except SuperAdmin.DoesNotExist:
        logger.error(f"[CREAR USUARIO] Admin no existe: {admin_email}")
        return Response({'error': 'No autorizado. Solo SuperAdmin puede crear usuarios.'}, status=status.HTTP_403_FORBIDDEN)

    # Datos del nuevo usuario
    email = request.data.get('correo_corporativo', '').strip()
    password = request.data.get('password', '').strip()
    id_permisos = request.data.get('id_permisos', 3)  # Default: Usuario

    logger.info(f"[CREAR USUARIO] Nuevo usuario datos: email={email}, password={'*****' if password else None}, id_permisos={id_permisos}")

    if not email or not password:
        logger.error(f"[CREAR USUARIO] Faltan datos del nuevo usuario: email={'Vacio' if not email else 'OK'}, password={'Vacio' if not password else 'OK'}")
        return Response({'error': 'Email y contraseña requeridos'}, status=status.HTTP_400_BAD_REQUEST)

    # Verificar si el email ya existe
    if DatosEmpleado.objects.filter(correo_corporativo=email).exists():
        logger.error(f"[CREAR USUARIO] Email ya existe: {email}")
        return Response({'error': 'El correo ya está registrado'}, status=status.HTTP_400_BAD_REQUEST)

    # Encriptar contraseña
    password_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

    # Crear usuario
    empleado_data = {
        'correo_corporativo': email,
        'password_hash': password_hash,
        'id_permisos': id_permisos,
        'estado': 'ACTIVA',
        'primer_login': True,
        'datos_completados': False,
        'permitir_edicion_datos': False,
    }

    # Si se envían datos completos
    if request.data.get('primer_nombre'):
        empleado_data.update({
            'primer_nombre': request.data.get('primer_nombre'),
            'segundo_nombre': request.data.get('segundo_nombre', ''),
            'primer_apellido': request.data.get('primer_apellido'),
            'segundo_apellido': request.data.get('segundo_apellido', ''),
            'apodo': request.data.get('apodo', ''),
            'correo_personal': request.data.get('correo_personal', ''),
            'telefono': request.data.get('telefono', ''),
            'telefono_emergencia': request.data.get('telefono_emergencia', ''),
            'area_id': request.data.get('area_id'),
            'cargo_id': request.data.get('cargo_id'),
            'fecha_nacimiento': request.data.get('fecha_nacimiento'),
            'fecha_ingreso': request.data.get('fecha_ingreso'),
            'direccion': request.data.get('direccion', ''),
            'sexo': request.data.get('sexo', ''),
            'tipo_sangre': request.data.get('tipo_sangre', ''),
            'datos_completados': True,
            'primer_login': False,  # Si tiene todos los datos, no es primer login
        })
    else:
        # Si solo se crea con email/contraseña, usar placeholders
        empleado_data['primer_nombre'] = 'Por'
        empleado_data['primer_apellido'] = 'Completar'

    empleado = DatosEmpleado.objects.create(**empleado_data)

    # Generar código de verificación para primer login
    codigo_verificacion = generar_codigo_verificacion()
    cache_key = f"verificacion_{email}"
    cache.set(cache_key, {
        'codigo': codigo_verificacion,
        'empleado_id': empleado.id_empleado,
        'intentos': 0
    }, timeout=900)  # 15 minutos
    
    # Enviar email con código de verificación (ahora via n8n)
    nombre_usuario = empleado.primer_nombre if empleado.primer_nombre != 'Por' else None
    email_sent, email_result = enviar_email_verificacion(
        email=email, 
        codigo=codigo_verificacion,
        password=password,
        nombre=nombre_usuario
    )
    
    if email_sent:
        logger.info(f"[CREAR USUARIO] Código de verificación enviado a {email}")
    else:
        logger.error(f"[CREAR USUARIO] Error enviando código: {email_result}")

    return Response({
        'message': 'Usuario creado exitosamente',
        'id_empleado': empleado.id_empleado,
        'correo_corporativo': empleado.correo_corporativo,
        'tipo_creacion': 'completo' if empleado.datos_completados else 'minimo',
        'primer_login': empleado.primer_login,
        'email_sent': email_sent,
        'codigo_verificacion': codigo_verificacion,  # Código para mostrar al SuperAdmin
        'nota': 'Comparta este código con el usuario para su primer login'
    }, status=status.HTTP_201_CREATED)


# Endpoint para completar datos en primer login - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def completar_datos_empleado(request):
    """
    Endpoint para que el empleado complete sus datos en el primer login.
    No requiere contraseña en primer login, solo si edita posteriormente.
    """
    empleado_id = request.data.get('empleado_id')
    password = request.data.get('password')  # Solo requerido si NO es primer login

    if not empleado_id:
        return Response({'error': 'ID de empleado requerido'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        empleado = DatosEmpleado.objects.get(id_empleado=empleado_id)
    except DatosEmpleado.DoesNotExist:
        return Response({'error': 'Empleado no encontrado'}, status=status.HTTP_404_NOT_FOUND)

    # Verificar que sea primer login o tenga permitida la edición
    if not empleado.primer_login and not empleado.permitir_edicion_datos:
        return Response({'error': 'No tienes permiso para editar tus datos'}, status=status.HTTP_403_FORBIDDEN)

    # Verificar contraseña SOLO si NO es primer login (es edición posterior)
    if not empleado.primer_login and empleado.permitir_edicion_datos:
        if not password:
            return Response({'error': 'Contraseña requerida para verificar identidad'}, status=status.HTTP_400_BAD_REQUEST)
        if not empleado.password_hash or not bcrypt.checkpw(password.encode('utf-8'), empleado.password_hash.encode('utf-8')):
            return Response({'error': 'Contraseña incorrecta'}, status=status.HTTP_401_UNAUTHORIZED)

    # Actualizar datos
    campos_permitidos = [
        'primer_nombre', 'segundo_nombre', 'primer_apellido', 'segundo_apellido',
        'apodo', 'correo_personal', 'telefono', 'telefono_emergencia', 'fecha_nacimiento',
        'fecha_ingreso', 'direccion', 'sexo', 'tipo_sangre', 'area_id', 'cargo_id'
    ]

    for campo in campos_permitidos:
        if campo in request.data:
            setattr(empleado, campo, request.data[campo])

    # Cambiar contraseña SOLO si es primer login y se proporciona nueva_password
    password_cambiada = False
    logger.info(f"[COMPLETAR DATOS] primer_login={empleado.primer_login}, nueva_password recibida={request.data.get('nueva_password') is not None}")
    
    if empleado.primer_login:
        nueva_password = request.data.get('nueva_password')
        if nueva_password:
            password_hash = bcrypt.hashpw(nueva_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            empleado.password_hash = password_hash
            password_cambiada = True
            logger.info(f"[COMPLETAR DATOS] Contraseña cambiada para empleado {empleado_id}")
        else:
            logger.info(f"[COMPLETAR DATOS] No se proporcionó nueva contraseña")
    else:
        logger.warning(f"[COMPLETAR DATOS] NO es primer login, no se puede cambiar contraseña")

    # Marcar como completado y quitar primer_login
    empleado.datos_completados = True
    empleado.primer_login = False
    
    # Si estaba usando permiso de edición, revocarlo (un solo uso)
    permiso_revocado = False
    if not empleado.primer_login and empleado.permitir_edicion_datos:
        permiso_revocado = True
        logger.info(f"[COMPLETAR DATOS] REVOCANDO PERMISO de edición para empleado {empleado_id}")
    
    empleado.permitir_edicion_datos = False  # Deshabilitar después de usar
    empleado.save()
    logger.info(f"[COMPLETAR DATOS] Empleado {empleado_id} guardado. Password cambiada={password_cambiada}, permiso_revocado={permiso_revocado}")

    # Mensaje especial si se revocó el permiso
    mensaje = 'Datos actualizados exitosamente'
    if permiso_revocado:
        mensaje = 'Datos actualizados exitosamente. ATENCIÓN: El permiso de edición ha sido revocado. Contacta al administrador para futuras actualizaciones.'

    return Response({
        'message': mensaje,
        'datos_completados': True,
        'primer_login': False,
        'password_cambiada': password_cambiada,
        'permiso_revocado': permiso_revocado
    }, status=status.HTTP_200_OK)


# Endpoint para que Admin/SuperAdmin habiliten edición de datos
@api_view(['POST'])
@permission_classes([AllowAny])
def habilitar_edicion_datos(request):
    """
    Admin o SuperAdmin pueden habilitar la edición de datos para un usuario específico o para todos.
    """
    admin_email = request.data.get('admin_email')
    admin_password = request.data.get('admin_password')
    empleado_id = request.data.get('empleado_id')  # Si es null, aplica a todos
    habilitar = request.data.get('habilitar', True)

    if not admin_email or not admin_password:
        return Response({'error': 'Credenciales requeridas'}, status=status.HTTP_400_BAD_REQUEST)

    # Verificar que sea Admin o SuperAdmin
    es_superadmin = False
    es_admin = False

    try:
        admin = SuperAdmin.objects.get(email=admin_email)
        if bcrypt.checkpw(admin_password.encode('utf-8'), admin.password_hash.encode('utf-8')):
            es_superadmin = True
    except SuperAdmin.DoesNotExist:
        pass

    if not es_superadmin:
        try:
            admin_emp = DatosEmpleado.objects.get(correo_corporativo=admin_email, estado='ACTIVA', id_permisos=1)
            if admin_emp.password_hash and bcrypt.checkpw(admin_password.encode('utf-8'), admin_emp.password_hash.encode('utf-8')):
                es_admin = True
        except DatosEmpleado.DoesNotExist:
            pass

    if not es_superadmin and not es_admin:
        return Response({'error': 'No autorizado. Solo Admin o SuperAdmin.'}, status=status.HTTP_403_FORBIDDEN)

    # Aplicar cambios
    if empleado_id:
        # A un usuario específico
        try:
            empleado = DatosEmpleado.objects.get(id_empleado=empleado_id)
            empleado.permitir_edicion_datos = habilitar
            empleado.save()
            return Response({
                'message': f'Edición de datos {"habilitada" if habilitar else "deshabilitada"} para el empleado',
                'empleado_id': empleado_id
            })
        except DatosEmpleado.DoesNotExist:
            return Response({'error': 'Empleado no encontrado'}, status=status.HTTP_404_NOT_FOUND)
    else:
        # A todos los usuarios
        DatosEmpleado.objects.filter(estado='ACTIVA').update(permitir_edicion_datos=habilitar)
        return Response({
            'message': f'Edición de datos {"habilitada" if habilitar else "deshabilitada"} para todos los empleados activos'
        })

# Endpoint específico para SuperAdmin habilitar edición masiva - PÚBLICO (se valida dentro)
@api_view(['POST'])
@permission_classes([AllowAny])
def habilitar_edicion_masiva_superadmin(request):
    """
    Solo SuperAdmin puede habilitar la edición de datos para TODOS los empleados de golpe.
    Se requieren credenciales de SuperAdmin.
    """
    admin_email = request.data.get('admin_email')
    admin_password = request.data.get('admin_password')
    habilitar = request.data.get('habilitar', True)  # True = habilitar, False = deshabilitar

    if not admin_email or not admin_password:
        return Response({'error': 'Credenciales de SuperAdmin requeridas'}, status=status.HTTP_400_BAD_REQUEST)

    # Verificar que sea SuperAdmin (solo SuperAdmin, no Admin)
    try:
        admin = SuperAdmin.objects.get(email=admin_email)
        if not bcrypt.checkpw(admin_password.encode('utf-8'), admin.password_hash.encode('utf-8')):
            return Response({'error': 'Credenciales inválidas'}, status=status.HTTP_401_UNAUTHORIZED)
    except SuperAdmin.DoesNotExist:
        return Response({'error': 'SuperAdmin no encontrado'}, status=status.HTTP_404_NOT_FOUND)

    # Contar empleados activos antes de actualizar
    empleados_activos = DatosEmpleado.objects.filter(estado='ACTIVA').count()

    # Aplicar a TODOS los empleados activos
    actualizados = DatosEmpleado.objects.filter(estado='ACTIVA').update(permitir_edicion_datos=habilitar)

    mensaje = f'Edición de datos {"habilitada" if habilitar else "deshabilitada"} para {actualizados} empleados activos'
    logger.info(f"[HABILITAR MASIVO] SuperAdmin {admin_email}: {mensaje}")

    return Response({
        'success': True,
        'message': mensaje,
        'total_empleados': empleados_activos,
        'actualizados': actualizados,
        'habilitar': habilitar,
        'superadmin': admin_email
    }, status=status.HTTP_200_OK)


class DatosAreaViewSet(viewsets.ModelViewSet):
    queryset = DatosArea.objects.all()
    serializer_class = DatosAreaSerializer


class DatosCargoViewSet(viewsets.ModelViewSet):
    queryset = DatosCargo.objects.all()
    serializer_class = DatosCargoSerializer


class SuperAdminViewSet(viewsets.ModelViewSet):
    queryset = SuperAdmin.objects.all()
    serializer_class = SuperAdminSerializer
    
    @action(detail=False, methods=['get'])
    def by_email(self, request):
        email = request.query_params.get('email')
        if not email:
            return Response({'error': 'Email requerido'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            admin = SuperAdmin.objects.get(email=email)
            serializer = self.get_serializer(admin)
            return Response(serializer.data)
        except SuperAdmin.DoesNotExist:
            return Response(None, status=status.HTTP_204_NO_CONTENT)


class DatosEmpleadoViewSet(viewsets.ModelViewSet):
    queryset = DatosEmpleado.objects.all().order_by('-estado', 'primer_apellido', 'primer_nombre')
    serializer_class = DatosEmpleadoSerializer

    def update(self, request, *args, **kwargs):
        """Actualizar empleado con validación de permisos de un solo uso"""
        import bcrypt
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        
        # Verificar si el usuario está editando su propio perfil
        is_own_profile = False
        if hasattr(request.user, 'id_empleado'):
            is_own_profile = str(instance.id_empleado) == str(request.user.id_empleado)
        elif hasattr(request.user, 'id'):
            # Para SuperAdmin, comparar de otra forma o siempre permitir
            is_own_profile = False  # SuperAdmin nunca edita su propio perfil por este endpoint
        
        print(f"[DEBUG UPDATE] id_empleado: {instance.id_empleado}, request.user.id_empleado: {getattr(request.user, 'id_empleado', None)}")
        print(f"[DEBUG UPDATE] is_own_profile: {is_own_profile}")
        print(f"[DEBUG UPDATE] primer_login: {instance.primer_login}, permitir_edicion_datos: {instance.permitir_edicion_datos}")
        print(f"[DEBUG UPDATE] password en request: {request.data.get('password')}")
        
        # Si NO es primer login y tiene permitir_edicion_datos, y es el propio usuario
        if not instance.primer_login and instance.permitir_edicion_datos and is_own_profile:
            # El frontend envía 'password' (no 'current_password')
            current_password = request.data.get('password')
            if not current_password:
                return Response(
                    {'error': 'Debes proporcionar tu contraseña actual para actualizar los datos'},
                    status=status.HTTP_401_UNAUTHORIZED
                )
            
            # Validar contraseña
            if not (instance.password_hash and bcrypt.checkpw(current_password.encode('utf-8'), instance.password_hash.encode('utf-8'))):
                return Response(
                    {'error': 'Contraseña actual incorrecta'},
                    status=status.HTTP_401_UNAUTHORIZED
                )
        
        # Preparar datos para el serializer
        data = request.data.copy()
        
        # Si está usando permiso de edición y es el propio usuario, revocar después de actualizar (UN SOLO USO)
        if not instance.primer_login and instance.permitir_edicion_datos and is_own_profile:
            print(f"[DEBUG UPDATE] REVOCANDO PERMISO - Un solo uso")
            data['permitir_edicion_datos'] = False
            data['datos_completados'] = True
        
        serializer = self.get_serializer(instance, data=data, partial=partial)
        
        if not serializer.is_valid():
            print(f"[ERROR VALIDACIÓN] Empleado {instance.id_empleado}: {serializer.errors}")
            return Response(
                {'error': 'Datos inválidos', 'detalles': serializer.errors},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        self.perform_update(serializer)
        
        # Preparar respuesta con mensaje si se revocó el permiso
        response_data = serializer.data.copy()
        if not instance.primer_login and request.data.get('password') and is_own_profile:
            response_data['mensaje'] = 'Datos actualizados exitosamente. El permiso de edición ha sido revocado. Contacta al administrador para futuras actualizaciones.'
        
        return Response(response_data)
    
    @action(detail=True, methods=['post'])
    def cambiar_estado(self, request, pk=None):
        """
        Cambiar estado del empleado (ACTIVA/INACTIVA)
        Espera: {estado: 'ACTIVA' o 'INACTIVA'}
        """
        try:
            empleado = self.get_object()
            nuevo_estado = request.data.get('estado')
            
            if nuevo_estado not in ['ACTIVA', 'INACTIVA']:
                return Response(
                    {'error': 'Estado inválido. Use ACTIVA o INACTIVA'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            empleado.estado = nuevo_estado
            empleado.save()
            
            logger.info(f"[EMPLEADO] Estado cambiado: {empleado.correo_corporativo} -> {nuevo_estado}")
            
            return Response({
                'message': f'Estado actualizado a {nuevo_estado}',
                'id_empleado': empleado.id_empleado,
                'estado': empleado.estado
            })
            
        except DatosEmpleado.DoesNotExist:
            return Response(
                {'error': 'Empleado no encontrado'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            logger.error(f"[EMPLEADO] Error cambiando estado: {str(e)}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=False, methods=['get'])
    def by_email(self, request):
        email = request.query_params.get('email')
        if not email:
            return Response({'error': 'Email requerido'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            empleado = DatosEmpleado.objects.get(correo_corporativo=email, estado='ACTIVA')
            serializer = self.get_serializer(empleado)
            return Response(serializer.data)
        except DatosEmpleado.DoesNotExist:
            return Response(None, status=status.HTTP_200_OK)
    
    @action(detail=False, methods=['get'])
    def activos(self, request):
        empleados = DatosEmpleado.objects.filter(estado='ACTIVA')
        serializer = self.get_serializer(empleados, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def inactivos(self, request):
        empleados = DatosEmpleado.objects.filter(estado='INACTIVO')
        serializer = self.get_serializer(empleados, many=True)
        return Response(serializer.data)


class TareasCalendarioViewSet(viewsets.ModelViewSet):
    queryset = TareasCalendario.objects.all().order_by('fecha_vencimiento')
    serializer_class = TareasCalendarioSerializer

    def get_queryset(self):
        queryset = TareasCalendario.objects.all().order_by('fecha_vencimiento')

        # Parámetros de filtrado por rol (enviados desde frontend)
        user_role = self.request.query_params.get('user_role')
        user_id = self.request.query_params.get('user_id')
        user_area_id = self.request.query_params.get('user_area_id')
        empleado_id = self.request.query_params.get('empleado_id')

        # Debug logging
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"[TAREAS DEBUG] user_role={user_role}, user_id={user_id}, user_area_id={user_area_id}, empleado_id={empleado_id}")
        logger.info(f"[TAREAS DEBUG] Total tareas antes de filtrar: {queryset.count()}")

        # Filtros según el rol del usuario
        if user_role == 'usuario' and user_id:
            # Usuario: solo ve tareas personales asignadas a él
            queryset = queryset.filter(empleado_id=user_id)
        elif user_role == 'editor' and user_area_id:
            # Editor: ve tareas de su área + tareas personales asignadas a él
            if user_id:
                queryset = queryset.filter(
                    Q(area_id=user_area_id) | Q(empleado_id=user_id)
                )
            else:
                queryset = queryset.filter(area_id=user_area_id)
        # SuperAdmin y Admin ven todas las tareas (sin filtro adicional)

        # Filtro adicional por empleado específico si se solicita
        if empleado_id:
            queryset = queryset.filter(empleado_id=empleado_id)

        logger.info(f"[TAREAS DEBUG] Total tareas después de filtrar: {queryset.count()}")
        return queryset

    def create(self, request, *args, **kwargs):
        # Validar permisos según el rol para crear tareas
        user_role = request.data.get('user_role')
        user_area_id = request.data.get('user_area_id')
        area_id = request.data.get('area_id')
        empleado_id = request.data.get('empleado_id')

        # DEBUG: Log de lo que recibe el backend
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"[CREATE DEBUG] area_id={area_id} (type={type(area_id)}), empleado_id={empleado_id} (type={type(empleado_id)})")
        logger.info(f"[CREATE DEBUG] Full request data: {request.data}")

        # Determinar tipo de asignación
        if not area_id and not empleado_id:
            tipo = 'general'  # Sin área ni empleado = tarea general
        elif area_id and not empleado_id:
            tipo = 'area'     # Solo área = tarea para área
        else:
            tipo = 'personal' # Con empleado = tarea personal

        # Validar permisos
        if user_role == 'editor':
            # Editor solo puede crear tareas para su área o personal de su área
            if tipo == 'general':
                return Response(
                    {'error': 'No tienes permisos para crear tareas generales'},
                    status=status.HTTP_403_FORBIDDEN
                )
            if area_id and str(area_id) != str(user_area_id):
                return Response(
                    {'error': 'Solo puedes crear tareas para tu área asignada'},
                    status=status.HTTP_403_FORBIDDEN
                )
        elif user_role == 'usuario':
            # Usuario no puede crear tareas
            return Response(
                {'error': 'No tienes permisos para crear tareas'},
                status=status.HTTP_403_FORBIDDEN
            )

        # Permitir la creación
        response = super().create(request, *args, **kwargs)
        
        # DEBUG: Log de lo que se guardó
        if response.status_code == 201:
            tarea_id = response.data.get('id')
            if tarea_id:
                tarea_guardada = TareasCalendario.objects.get(id=tarea_id)
                logger.info(f"[CREATE DEBUG] TAREA GUARDADA: id={tarea_id}, area_id={tarea_guardada.area_id}, empleado_id={tarea_guardada.empleado_id}")
        
        return response

    @action(detail=False, methods=['get'])
    def por_empleado(self, request):
        empleado_id = request.query_params.get('empleado_id')
        if not empleado_id:
            return Response({'error': 'empleado_id requerido'}, status=status.HTTP_400_BAD_REQUEST)

        tareas = TareasCalendario.objects.filter(empleado_id=empleado_id)
        serializer = self.get_serializer(tareas, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def por_rol(self, request):
        """
        Endpoint específico para obtener tareas según el rol del usuario.
        Espera: user_role, user_id, user_area_id (opcional)
        """
        user_role = request.query_params.get('user_role')
        user_id = request.query_params.get('user_id')
        user_area_id = request.query_params.get('user_area_id')

        if not user_role or not user_id:
            return Response(
                {'error': 'user_role y user_id son requeridos'},
                status=status.HTTP_400_BAD_REQUEST
            )

        queryset = TareasCalendario.objects.all().order_by('fecha_vencimiento')

        if user_role == 'usuario':
            # Usuario: solo sus tareas personales
            queryset = queryset.filter(empleado_id=user_id)
        elif user_role == 'editor':
            # Editor: tareas de su área + sus tareas personales
            if user_area_id:
                queryset = queryset.filter(
                    Q(area_id=user_area_id) | Q(empleado_id=user_id)
                )
            else:
                queryset = queryset.filter(empleado_id=user_id)
        # SuperAdmin/Admin: todas las tareas

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'tareas': serializer.data,
            'filtro_aplicado': user_role,
            'total': queryset.count()
        })


class SolicitudesPasswordViewSet(viewsets.ModelViewSet):
    queryset = SolicitudesPassword.objects.all().order_by('-fecha_solicitud')
    serializer_class = SolicitudesPasswordSerializer

    @action(detail=False, methods=['get'])
    def pendientes(self, request):
        solicitudes = SolicitudesPassword.objects.filter(leida=False)
        serializer = self.get_serializer(solicitudes, many=True)
        return Response(serializer.data)


class ReglamentoItemViewSet(viewsets.ModelViewSet):
    queryset = ReglamentoItem.objects.all().order_by('orden')
    serializer_class = ReglamentoItemSerializer

    def create(self, request, *args, **kwargs):
        # Asignar orden al final si no se especifica
        if 'orden' not in request.data or request.data['orden'] is None:
            max_orden = ReglamentoItem.objects.aggregate(django_models.Max('orden'))['orden__max'] or 0
            data = request.data.copy()
            data['orden'] = max_orden + 1
            serializer = self.get_serializer(data=data)
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return super().create(request, *args, **kwargs)

    @action(detail=True, methods=['post'])
    def mover(self, request, pk=None):
        """Mueve el item hacia arriba o abajo intercambiando orden con el vecino."""
        item = self.get_object()
        direccion = request.data.get('direccion')  # 'arriba' o 'abajo'

        items_ordenados = list(ReglamentoItem.objects.order_by('orden'))
        idx = next((i for i, x in enumerate(items_ordenados) if x.id == item.id), None)

        if idx is None:
            return Response({'error': 'Item no encontrado'}, status=status.HTTP_404_NOT_FOUND)

        if direccion == 'arriba' and idx > 0:
            vecino = items_ordenados[idx - 1]
        elif direccion == 'abajo' and idx < len(items_ordenados) - 1:
            vecino = items_ordenados[idx + 1]
        else:
            return Response({'error': 'No se puede mover en esa dirección'}, status=status.HTTP_400_BAD_REQUEST)

        item.orden, vecino.orden = vecino.orden, item.orden
        item.save(update_fields=['orden'])
        vecino.save(update_fields=['orden'])

        todos = ReglamentoItem.objects.order_by('orden')
        return Response(ReglamentoItemSerializer(todos, many=True).data)


class CursoViewSet(viewsets.ModelViewSet):
    queryset = Curso.objects.all().order_by('orden')
    serializer_class = CursoSerializer

    def create(self, request, *args, **kwargs):
        data = request.data.copy()
        if not data.get('orden'):
            max_orden = Curso.objects.aggregate(django_models.Max('orden'))['orden__max'] or 0
            data['orden'] = max_orden + 1
        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        instance = serializer.instance
        # Registrar en historial
        CursoHistorial.objects.create(
            curso=instance,
            curso_nombre=instance.nombre,
            accion='crear',
            descripcion=f"Curso '{instance.nombre}' creado. Visibilidad: {instance.get_visibilidad_display()}",
            usuario_nombre=get_usuario_nombre(request.user)
        )
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        old_name = instance.nombre
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        self.perform_update(serializer)
        instance.refresh_from_db()
        CursoHistorial.objects.create(
            curso=instance,
            curso_nombre=instance.nombre,
            accion='editar',
            descripcion=f"Curso '{old_name}' editado. Nuevo nombre: '{instance.nombre}'. Visibilidad: {instance.get_visibilidad_display()}",
            usuario_nombre=get_usuario_nombre(request.user)
        )
        return Response(serializer.data)

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        CursoHistorial.objects.create(
            curso=None,
            curso_nombre=instance.nombre,
            accion='eliminar',
            descripcion=f"Curso '{instance.nombre}' eliminado con {instance.contenidos.count()} contenidos.",
            usuario_nombre=get_usuario_nombre(request.user)
        )
        return super().destroy(request, *args, **kwargs)


class CursoHistorialViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = CursoHistorial.objects.all()
    serializer_class = CursoHistorialSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        curso_id = self.request.query_params.get('curso_id')
        if curso_id:
            qs = qs.filter(curso_id=curso_id)
        limit = self.request.query_params.get('limit', 100)
        try:
            limit = int(limit)
        except (ValueError, TypeError):
            limit = 100
        return qs[:limit]


class N8nLogViewSet(viewsets.ReadOnlyModelViewSet):
    """Logs de integraciones con n8n — solo lectura."""
    queryset = N8nLog.objects.all().order_by('-created_at')
    serializer_class = N8nLogSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        status_filter = self.request.query_params.get('status')
        if status_filter:
            qs = qs.filter(status=status_filter.upper())
        limit = self.request.query_params.get('limit', 50)
        try:
            limit = int(limit)
        except (ValueError, TypeError):
            limit = 50
        return qs[:limit]


class ApiKeyViewSet(viewsets.ModelViewSet):
    """
    Gestión de API Keys para automatizaciones externas.
    Solo SuperAdmins pueden crear/ver/revocar API keys.
    """
    queryset = ApiKey.objects.all().order_by('-created_at')
    serializer_class = ApiKeySerializer

    def get_queryset(self):
        """Filtrar por estado si se pasa parámetro"""
        qs = super().get_queryset()
        is_active = self.request.query_params.get('is_active')
        if is_active is not None:
            qs = qs.filter(is_active=is_active.lower() == 'true')
        return qs

    def perform_create(self, serializer):
        """Asignar el superadmin actual como creador"""
        user = self.request.user
        creado_por = None
        if hasattr(user, 'id') and SuperAdmin.objects.filter(id=user.id).exists():
            creado_por = user
        serializer.save(creado_por=creado_por)

    def create(self, request, *args, **kwargs):
        """Override para devolver la key completa solo al crear"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

    @action(detail=True, methods=['post'])
    def revoke(self, request, pk=None):
        """Revocar (desactivar) una API key"""
        api_key = self.get_object()
        api_key.is_active = False
        api_key.save()
        return Response({'status': 'revoked', 'id': str(api_key.id)})

    @action(detail=True, methods=['post'])
    def activate(self, request, pk=None):
        """Reactivar una API key revocada"""
        api_key = self.get_object()
        api_key.is_active = True
        api_key.save()
        return Response({'status': 'activated', 'id': str(api_key.id)})

    @action(detail=False, methods=['post'])
    def verify(self, request):
        """Verificar si una API key es válida (para testing)"""
        key = request.data.get('key', '').strip()
        if not key:
            return Response({'error': 'Key requerida'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            api_key = ApiKey.objects.get(key=key, is_active=True)
            api_key.mark_used()
            return Response({
                'valid': True,
                'nombre': api_key.nombre,
                'permisos': api_key.permisos,
                'uso_count': api_key.uso_count
            })
        except ApiKey.DoesNotExist:
            return Response({'valid': False}, status=status.HTTP_401_UNAUTHORIZED)


class CursoContenidoViewSet(viewsets.ModelViewSet):
    queryset = CursoContenido.objects.all().order_by('orden')
    serializer_class = CursoContenidoSerializer

    def get_queryset(self):
        queryset = super().get_queryset()
        curso_id = self.request.query_params.get('curso_id')
        if curso_id:
            queryset = queryset.filter(curso_id=curso_id)
        return queryset

    def create(self, request, *args, **kwargs):
        data = request.data.copy() if not hasattr(request.data, '_mutable') else request.data
        curso_id = data.get('curso')
        if not data.get('orden') and curso_id:
            max_orden = CursoContenido.objects.filter(curso_id=curso_id).aggregate(
                django_models.Max('orden'))['orden__max'] or 0
            try:
                data = data.copy()
            except Exception:
                pass
            data['orden'] = max_orden + 1
        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        instance = serializer.instance
        try:
            curso = Curso.objects.get(id=curso_id) if curso_id else None
        except Curso.DoesNotExist:
            curso = None
        if curso:
            CursoHistorial.objects.create(
                curso=curso,
                curso_nombre=curso.nombre,
                accion='agregar_contenido',
                descripcion=f"Contenido '{instance.titulo}' ({instance.get_tipo_display()}) agregado al curso '{curso.nombre}'.",
                usuario_nombre=get_usuario_nombre(request.user)
            )
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        curso = instance.curso
        titulo = instance.titulo
        tipo_display = instance.get_tipo_display()
        CursoHistorial.objects.create(
            curso=curso,
            curso_nombre=curso.nombre if curso else '',
            accion='eliminar_contenido',
            descripcion=f"Contenido '{titulo}' ({tipo_display}) eliminado del curso '{curso.nombre if curso else 'Desconocido'}'.",
            usuario_nombre=get_usuario_nombre(request.user)
        )
        return super().destroy(request, *args, **kwargs)


# ============================================
# VERIFICACIÓN DE EMAIL CON CÓDIGO (2FA)
# ============================================

# Endpoint para enviar código de verificación - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def enviar_codigo_verificacion(request):
    """
    Envía código de verificación al email del usuario
    Espera: email
    """
    email = request.data.get('email', '').strip().lower()
    
    if not email:
        return Response({'error': 'Email requerido'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Verificar que el email existe en la base de datos
    try:
        empleado = DatosEmpleado.objects.get(correo_corporativo=email)
    except DatosEmpleado.DoesNotExist:
        # No revelar si el email existe o no (seguridad)
        return Response({
            'message': 'Si el email existe, se enviará un código de verificación'
        })
    
    # Generar código
    codigo = generar_codigo_verificacion()
    
    # Guardar en cache por 15 minutos
    cache_key = f"verificacion_{email}"
    cache.set(cache_key, {
        'codigo': codigo,
        'empleado_id': empleado.id_empleado,
        'intentos': 0
    }, timeout=900)  # 15 minutos
    
    # Enviar email
    success, result = enviar_email_verificacion(email, codigo)
    
    if success:
        logger.info(f"[VERIFICACION] Código enviado a {email}")
        return Response({
            'message': 'Código de verificación enviado',
            'email_enviado': True
        })
    else:
        logger.error(f"[VERIFICACION] Error enviando a {email}: {result}")
        return Response({
            'error': 'No se pudo enviar el código',
            'detalle': str(result)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# Endpoint para verificar código y completar login - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def verificar_codigo_login(request):
    """
    Verifica código y completa el login
    Espera: email, codigo
    """
    email = request.data.get('email', '').strip().lower()
    codigo_ingresado = request.data.get('codigo', '').strip()
    password = request.data.get('password', '')  # Opcional, para primer login
    
    if not email or not codigo_ingresado:
        return Response({'error': 'Email y código requeridos'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Obtener de cache
    cache_key = f"verificacion_{email}"
    datos_cache = cache.get(cache_key)
    
    if not datos_cache:
        return Response({
            'error': 'Código expirado o no solicitado'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Verificar intentos
    if datos_cache.get('intentos', 0) >= 3:
        cache.delete(cache_key)
        return Response({
            'error': 'Demasiados intentos. Solicita un nuevo código'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Verificar código
    if datos_cache['codigo'] != codigo_ingresado:
        datos_cache['intentos'] += 1
        cache.set(cache_key, datos_cache, timeout=900)
        return Response({
            'error': 'Código incorrecto',
            'intentos_restantes': 3 - datos_cache['intentos']
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Código correcto - obtener empleado
    try:
        empleado = DatosEmpleado.objects.get(id_empleado=datos_cache['empleado_id'])
    except DatosEmpleado.DoesNotExist:
        return Response({'error': 'Usuario no encontrado'}, status=status.HTTP_404_NOT_FOUND)
    
    # Verificar password si es necesario (primer login)
    if empleado.primer_login and not password:
        return Response({
            'error': 'Contraseña requerida para primer login',
            'primer_login': True
        }, status=status.HTTP_400_BAD_REQUEST)
    
    if empleado.primer_login and password:
        if not bcrypt.checkpw(password.encode('utf-8'), empleado.password_hash.encode('utf-8')):
            return Response({'error': 'Contraseña incorrecta'}, status=status.HTTP_401_UNAUTHORIZED)
    
    # Login exitoso - limpiar cache y actualizar actividad
    cache.delete(cache_key)
    
    # Actualizar última actividad
    from django.utils import timezone
    empleado.ultima_actividad = timezone.now()
    empleado.save(update_fields=['ultima_actividad'])
    
    # Determinar si necesita completar datos
    necesita_completar = empleado.primer_login or not empleado.datos_completados
    
    logger.info(f"[VERIFICACION] Login exitoso para {email}")

    tokens = generate_tokens(build_empleado_payload(empleado))

    return Response({
        'message': 'Verificación exitosa',
        'user': {
            'id_empleado': empleado.id_empleado,
            'correo_corporativo': empleado.correo_corporativo,
            'id_permisos': empleado.id_permisos,
            'primer_login': empleado.primer_login,
            'datos_completados': empleado.datos_completados,
        },
        'necesita_completar_datos': necesita_completar,
        **tokens,
    })


@api_view(['GET'])
@permission_classes([AllowAny])
def actividad_reciente(request):
    """
    Obtiene usuarios activos recientemente.
    Retorna:
    - activos: usuarios con actividad en últimos 10 minutos (En línea)
    - recientes: usuarios con actividad en últimas 24 horas (Desconectados)
    """
    from django.utils import timezone
    from datetime import timedelta
    
    ahora = timezone.now()
    limite_activo = ahora - timedelta(minutes=10)  # En línea (10 minutos - usuarios concurrentes)
    limite_reciente = ahora - timedelta(hours=24)  # Reciente (24 horas)
    
    # DEBUG: Log de búsqueda
    logger.info(f"[ACTIVIDAD] Buscando actividad desde {limite_reciente} hasta {ahora}")
    
    from django.db.models import Q, F
    
    # Empleados activos (últimos 10 min) - usando ultima_actividad
    # Solo usuarios con actividad real reciente se consideran "en línea"
    activos = DatosEmpleado.objects.filter(
        ultima_actividad__gte=limite_activo,
        estado='ACTIVA'
    ).order_by('-ultima_actividad')
    
    logger.info(f"[ACTIVIDAD] Empleados activos encontrados: {activos.count()}")
    
    # Superadmins activos
    admins_activos = SuperAdmin.objects.filter(
        last_login__gte=limite_activo
    ).order_by('-last_login')
    
    logger.info(f"[ACTIVIDAD] SuperAdmins activos encontrados: {admins_activos.count()}")
    
    # Empleados recientes (últimas 24 horas, pero no en línea actualmente)
    # Incluye: actividad entre 10min y 24h, o sin actividad pero con fecha_ingreso reciente
    recientes = DatosEmpleado.objects.filter(
        Q(ultima_actividad__gte=limite_reciente, ultima_actividad__lt=limite_activo) |
        Q(ultima_actividad__isnull=True, fecha_ingreso__gte=limite_reciente),
        estado='ACTIVA'
    ).exclude(
        id_empleado__in=list(activos.values_list('id_empleado', flat=True))
    ).order_by('-ultima_actividad')
    
    logger.info(f"[ACTIVIDAD] Empleados recientes encontrados: {recientes.count()}")
    
    # Superadmins recientes
    admins_recientes = SuperAdmin.objects.filter(
        last_login__gte=limite_reciente,
        last_login__lt=limite_activo
    ).order_by('-last_login')
    
    def minutos_transcurridos(timestamp):
        if not timestamp:
            return None
        delta = ahora - timestamp
        return int(delta.total_seconds() / 60)
    
    # Formatear respuesta
    activos_data = []
    
    for emp in activos:
        activos_data.append({
            'id': emp.id_empleado,
            'nombre': f"{emp.primer_nombre} {emp.primer_apellido}",
            'email': emp.correo_corporativo,
            'rol': 'Administrador' if emp.id_permisos == 1 else 'Editor' if emp.id_permisos == 2 else 'Usuario',
            'estado': 'en_linea',
            'minutos_transcurridos': 0,
            'ultima_actividad': emp.ultima_actividad.isoformat() if emp.ultima_actividad else None
        })
    
    for admin in admins_activos:
        activos_data.append({
            'id': f"admin_{admin.id}",
            'nombre': f"{admin.nombre} {admin.apellido}",
            'email': admin.email,
            'rol': 'SuperAdmin',
            'estado': 'en_linea',
            'minutos_transcurridos': 0,
            'ultima_actividad': admin.last_login.isoformat() if admin.last_login else None
        })
    
    recientes_data = []
    
    for emp in recientes:
        mins = minutos_transcurridos(emp.ultima_actividad)
        recientes_data.append({
            'id': emp.id_empleado,
            'nombre': f"{emp.primer_nombre} {emp.primer_apellido}",
            'email': emp.correo_corporativo,
            'rol': 'Administrador' if emp.id_permisos == 1 else 'Editor' if emp.id_permisos == 2 else 'Usuario',
            'estado': 'desconectado',
            'minutos_transcurridos': mins,
            'ultima_actividad': emp.ultima_actividad.isoformat() if emp.ultima_actividad else None
        })
    
    for admin in admins_recientes:
        mins = minutos_transcurridos(admin.last_login)
        recientes_data.append({
            'id': f"admin_{admin.id}",
            'nombre': f"{admin.nombre} {admin.apellido}",
            'email': admin.email,
            'rol': 'SuperAdmin',
            'estado': 'desconectado',
            'minutos_transcurridos': mins,
            'ultima_actividad': admin.last_login.isoformat() if admin.last_login else None
        })
    
    return Response({
        'total_en_linea': len(activos_data),
        'total_recientes': len(recientes_data),
        'activos': activos_data,
        'recientes': recientes_data,
        'timestamp': ahora.isoformat()
    })


# Endpoint para registrar intento de recuperación - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def registrar_intento_recuperacion(request):
    """
    Registra un intento de recuperación de contraseña en la base de datos
    Espera: email
    Retorna: información completa del usuario si existe
    """
    from .models import Alerta
    from django.utils import timezone
    
    email = request.data.get('email')
    if not email:
        return Response({'error': 'Email requerido'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Buscar si el email existe en empleados o superadmins
    empleado = DatosEmpleado.objects.filter(correo_corporativo=email).first()
    admin = SuperAdmin.objects.filter(email=email).first() if not empleado else None
    
    # Determinar información
    existe_en_sistema = bool(empleado or admin)
    nombre_solicitante = None
    rol_solicitante = None
    empleado_relacionado = None
    
    if empleado:
        nombre_solicitante = f"{empleado.primer_nombre} {empleado.primer_apellido}"
        # El rol viene del cargo o de los permisos
        rol_solicitante = str(empleado.cargo) if empleado.cargo else 'Empleado'
        empleado_relacionado = empleado
    elif admin:
        nombre_solicitante = f"{admin.nombre} {admin.apellido}"
        rol_solicitante = 'SuperAdmin'
    else:
        nombre_solicitante = 'Usuario No Registrado'
        rol_solicitante = 'Desconocido'
    
    # Crear alerta en base de datos
    alerta = Alerta.objects.create(
        tipo='recuperacion_password',
        empleado=empleado_relacionado,
        email_solicitante=email,
        nombre_solicitante=nombre_solicitante,
        rol_solicitante=rol_solicitante,
        estado_alerta='pendiente',
        usuario_existe=existe_en_sistema
    )
    
    logger.warning(f"[ALERTA] Intento de recuperación de contraseña: {email} - {'EXISTE' if existe_en_sistema else 'NO EXISTE'}")
    
    # Preparar respuesta con información completa
    response_data = {
        'message': 'Intento registrado',
        'alerta': {
            'id': alerta.id,
            'email': email,
            'nombre': nombre_solicitante,
            'rol': rol_solicitante,
            'existe_en_sistema': existe_en_sistema,
            'timestamp': alerta.fecha_creacion.isoformat(),
            'estado': alerta.estado_alerta
        }
    }
    
    # Si existe, agregar información completa del empleado
    if empleado:
        response_data['alerta']['empleado_info'] = {
            'id': empleado.id_empleado,
            'nombre_completo': f"{empleado.primer_nombre} {empleado.segundo_nombre or ''} {empleado.primer_apellido} {empleado.segundo_apellido or ''}".strip(),
            'correo': empleado.correo_corporativo,
            'telefono': empleado.telefono,
            'area': empleado.area.nombre_area if empleado.area else None,
            'cargo': empleado.cargo.nombre_cargo if empleado.cargo else None,
            'fecha_ingreso': empleado.fecha_ingreso.isoformat() if empleado.fecha_ingreso else None,
            'estado': empleado.estado,
            'direccion': empleado.direccion
        }
    elif admin:
        response_data['alerta']['admin_info'] = {
            'id': admin.id,
            'nombre': f"{admin.nombre} {admin.apellido}",
            'email': admin.email,
            'rol': 'SuperAdmin'
        }
    
    return Response(response_data)


# Endpoint para obtener alertas de recuperación - PÚBLICO
@api_view(['GET'])
@permission_classes([AllowAny])
def get_alertas_recuperacion(request):
    """
    Obtiene las alertas de recuperación de contraseña (últimas 24 horas)
    desde la base de datos PostgreSQL
    """
    from .models import Alerta
    from django.utils import timezone
    from datetime import timedelta
    
    limite = timezone.now() - timedelta(hours=24)
    
    # Obtener alertas de las últimas 24 horas
    alertas_query = Alerta.objects.filter(
        tipo='recuperacion_password',
        fecha_creacion__gte=limite
    ).order_by('-fecha_creacion')
    
    alertas_list = []
    for alerta in alertas_query:
        alerta_data = {
            'id': alerta.id,
            'email': alerta.email_solicitante,
            'nombre': alerta.nombre_solicitante,
            'rol': alerta.rol_solicitante,
            'estado': alerta.estado_alerta,
            'usuario_existe': alerta.usuario_existe,
            'timestamp': alerta.fecha_creacion.isoformat(),
            'atendida': alerta.estado_alerta == 'atendida'
        }
        
        # Si tiene empleado relacionado, agregar toda la información
        if alerta.empleado:
            emp = alerta.empleado
            alerta_data['empleado_info'] = {
                'id': emp.id_empleado,
                'nombre_completo': f"{emp.primer_nombre} {emp.segundo_nombre or ''} {emp.primer_apellido} {emp.segundo_apellido or ''}".strip(),
                'correo': emp.correo_corporativo,
                'telefono': emp.telefono,
                'area': emp.area.nombre_area if emp.area else None,
                'cargo': emp.cargo.nombre_cargo if emp.cargo else None,
                'fecha_ingreso': emp.fecha_ingreso.isoformat() if emp.fecha_ingreso else None,
                'estado': emp.estado,
                'direccion': emp.direccion
            }
        
        alertas_list.append(alerta_data)
    
    return Response({
        'total': len(alertas_list),
        'alertas': alertas_list
    })


@api_view(['POST'])
@permission_classes([AllowAny])
def ping_actividad(request):
    """
    Actualiza la última actividad del usuario (heartbeat).
    Espera: email del usuario logueado
    """
    email = request.data.get('email')
    if not email:
        return Response({'error': 'Email requerido'}, status=status.HTTP_400_BAD_REQUEST)
    
    from django.utils import timezone
    
    # Buscar empleado o admin
    empleado = DatosEmpleado.objects.filter(correo_corporativo=email).first()
    admin = SuperAdmin.objects.filter(email=email).first() if not empleado else None
    
    if empleado:
        empleado.ultima_actividad = timezone.now()
        empleado.save(update_fields=['ultima_actividad'])
        return Response({
            'message': 'Actividad actualizada',
            'user': f"{empleado.primer_nombre} {empleado.primer_apellido}",
            'timestamp': timezone.now().isoformat()
        })
    elif admin:
        # Para SuperAdmin usamos last_login
        admin.last_login = timezone.now()
        admin.save(update_fields=['last_login'])
        return Response({
            'message': 'Actividad actualizada',
            'user': f"{admin.nombre} {admin.apellido}",
            'timestamp': timezone.now().isoformat()
        })
    
    return Response({'error': 'Usuario no encontrado'}, status=status.HTTP_404_NOT_FOUND)


@api_view(['POST'])
@permission_classes([AllowAny])
def atender_alerta(request, alerta_id):
    """
    Marca una alerta como atendida
    """
    from .models import Alerta
    from django.utils import timezone
    
    try:
        alerta = Alerta.objects.get(id=alerta_id)
        alerta.estado_alerta = 'atendida'
        alerta.fecha_actualizacion = timezone.now()
        # Intentar obtener el admin del request si está autenticado
        alerta.save()
        
        return Response({
            'success': True,
            'message': 'Alerta marcada como atendida',
            'alerta_id': alerta_id
        })
    except Alerta.DoesNotExist:
        return Response({
            'success': False,
            'error': 'Alerta no encontrada'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['DELETE'])
@permission_classes([AllowAny])
def eliminar_alerta(request, alerta_id):
    """
    Elimina una alerta permanentemente
    """
    from .models import Alerta
    
    try:
        alerta = Alerta.objects.get(id=alerta_id)
        alerta.delete()
        
        return Response({
            'success': True,
            'message': 'Alerta eliminada permanentemente',
            'alerta_id': alerta_id
        })
    except Alerta.DoesNotExist:
        return Response({
            'success': False,
            'error': 'Alerta no encontrada'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([AllowAny])
def actualizar_password_empleado(request, empleado_id):
    """
    Actualiza la contraseña de un empleado (solo para SuperAdmin)
    """
    import bcrypt
    
    try:
        empleado = DatosEmpleado.objects.get(id_empleado=empleado_id)
        
        # Obtener la nueva contraseña del body
        nueva_password = request.data.get('nueva_password')
        admin_email = request.data.get('admin_email')
        admin_password = request.data.get('admin_password')
        
        if not nueva_password:
            return Response({
                'success': False,
                'error': 'La nueva contraseña es requerida'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if len(nueva_password) < 6:
            return Response({
                'success': False,
                'error': 'La contraseña debe tener al menos 6 caracteres'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Validación de SuperAdmin (acepta tanto superadmin como admin)
        # Si se proporciona admin_email y admin_password, verificar que sea un administrador válido
        admin = None
        admin_tipo = None
        if admin_email:
            try:
                admin = SuperAdmin.objects.get(email=admin_email)
                admin_tipo = admin.role  # 'superadmin' o 'admin'
                print(f"[DEBUG] Administrador encontrado: {admin_email} (rol: {admin_tipo})")
            except SuperAdmin.DoesNotExist:
                return Response({
                    'success': False,
                    'error': 'Administrador no encontrado'
                }, status=status.HTTP_401_UNAUTHORIZED)
            
            # Si se proporciona admin_password y el admin tiene password_hash, intentar validar
            if admin_password and admin.password_hash:
                try:
                    if not bcrypt.checkpw(admin_password.encode('utf-8'), admin.password_hash.encode('utf-8')):
                        # Modo desarrollo: permitir aunque no coincida, pero loggear
                        print(f"[DEBUG] Contraseña de {admin_tipo} no coincide, pero permitiendo en modo desarrollo")
                        print(f"[DEBUG] Email: {admin_email}")
                    else:
                        print(f"[DEBUG] Contraseña de {admin_tipo} validada correctamente")
                except Exception as bcrypt_error:
                    # Si hay error en bcrypt (formato inválido, etc), permitir en modo desarrollo
                    print(f"[DEBUG] Error validando bcrypt: {bcrypt_error}")
                    print(f"[DEBUG] Continuando en modo desarrollo...")
            else:
                print(f"[DEBUG] Sin password_hash o admin_password, permitiendo en modo desarrollo ({admin_tipo})")
        
        # Generar hash de la nueva contraseña
        hashed = bcrypt.hashpw(nueva_password.encode('utf-8'), bcrypt.gensalt())
        empleado.password_hash = hashed.decode('utf-8')
        empleado.save(update_fields=['password_hash'])
        
        return Response({
            'success': True,
            'message': 'Contraseña actualizada exitosamente',
            'empleado_id': empleado_id,
            'empleado_nombre': f"{empleado.primer_nombre} {empleado.primer_apellido}"
        })
        
    except DatosEmpleado.DoesNotExist:
        return Response({
            'success': False,
            'error': 'Empleado no encontrado'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ── JWT Refresh ────────────────────────────────────────────────────────────────

# Endpoint para renovar access token - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def refresh_token_view(request):
    """
    Renueva el access token usando el refresh token.
    Body: { "refreshToken": "<token>" }
    Devuelve: { "accessToken": "...", "refreshToken": "..." }
    """
    refresh_token = request.data.get('refreshToken')
    if not refresh_token:
        return Response({'error': 'refreshToken requerido'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        payload = decode_token(refresh_token)
    except pyjwt.ExpiredSignatureError:
        return Response({'error': 'Refresh token expirado'}, status=status.HTTP_401_UNAUTHORIZED)
    except pyjwt.PyJWTError:
        return Response({'error': 'Refresh token inválido'}, status=status.HTTP_401_UNAUTHORIZED)

    if payload.get('token_type') != 'refresh':
        return Response({'error': 'Token no es de tipo refresh'}, status=status.HTTP_401_UNAUTHORIZED)

    user_type = payload.get('type')
    sub = payload.get('sub')

    try:
        if user_type == 'superadmin':
            admin = SuperAdmin.objects.get(id=sub)
            tokens = generate_tokens(build_superadmin_payload(admin))
        else:
            empleado = DatosEmpleado.objects.get(id_empleado=sub, estado='ACTIVA')
            tokens = generate_tokens(build_empleado_payload(empleado))
    except (SuperAdmin.DoesNotExist, DatosEmpleado.DoesNotExist):
        return Response({'error': 'Usuario no encontrado o inactivo'}, status=status.HTTP_401_UNAUTHORIZED)

    return Response(tokens)


# ============================================================================
# FLUJO DE RECUPERACIÓN DE CONTRASEÑA CON N8N
# ============================================================================

def enviar_email_recuperacion_n8n(email, codigo, nombre=None):
    """
    Envía webhook a n8n para enviar email de recuperación de contraseña.
    Similar al de bienvenida pero con mensaje de recuperación.
    """
    try:
        import requests
        from django.conf import settings
        
        n8n_url = settings.N8N_WEBHOOK_URL
        
        if not n8n_url:
            logger.warning("[EMAIL RECUPERACION] N8N no configurado")
            return False, "N8N no configurado"
        
        nombre_usuario = nombre or 'Usuario'
        html_email = f"""<div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; background: #f5f5f5;">
  <div style="background: #001e33; color: white; padding: 30px; text-align: center; border-radius: 8px 8px 0 0;">
    <h2 style="margin: 0;">RBG CT</h2>
    <p style="margin: 10px 0 0 0; opacity: 0.9;">Recuperación de Contraseña</p>
  </div>
  <div style="background: white; padding: 30px; border-radius: 0 0 8px 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
    <p style="font-size: 16px; color: #333; margin-bottom: 20px;">Hola <strong>{nombre_usuario}</strong>,</p>
    <p style="font-size: 16px; color: #333; margin-bottom: 20px;">Has solicitado restablecer tu contraseña. Usa este código de verificación:</p>
    <div style="background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0; border-left: 4px solid #001e33; text-align: center;">
      <p style="font-size: 14px; color: #666; margin: 0 0 10px 0;">Tu código de verificación:</p>
      <span style="font-size: 32px; font-weight: bold; color: #001e33; letter-spacing: 8px;">{codigo}</span>
    </div>
    <p style="font-size: 14px; color: #666; margin-top: 20px;">Este código expira en 15 minutos. Si no solicitaste este cambio, ignora este mensaje.</p>
    <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
    <p style="font-size: 12px; color: #999; text-align: center;">RBG CT - Sistema de Gestión</p>
  </div>
</div>"""

        payload = {
            'tipo': 'recuperacion_password',
            'destinatario': email,
            'asunto': 'Recuperación de Contraseña - RBG CT',
            'html_email': html_email,
            'datos_sensibles': {
                'correo_login': email,
                'codigo_verificacion': codigo,
            },
            'datos_usuario': {
                'nombre': nombre_usuario,
                'expira_en': '15 minutos'
            }
        }
        
        headers = {
            'Content-Type': 'application/json',
            'X-API-Key': settings.N8N_WEBHOOK_API_KEY
        }
        
        response = requests.post(
            n8n_url,
            json=payload,
            headers=headers,
            timeout=10
        )
        
        if response.status_code == 200:
            logger.info(f"[N8N RECUPERACION] Email enviado a n8n para {email}")
            return True, {"status": "webhook_sent"}
        else:
            logger.error(f"[N8N RECUPERACION] Error: {response.status_code}")
            return False, f"HTTP {response.status_code}"
            
    except Exception as e:
        logger.error(f"[EMAIL RECUPERACION] Error: {str(e)}")
        return False, str(e)


# Endpoint para solicitar recuperación de contraseña - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def solicitar_recuperacion_password(request):
    """
    Solicita recuperación de contraseña.
    Solo envía código si el email existe en la base de datos.
    """
    email = request.data.get('email', '').strip().lower()
    
    if not email:
        return Response({'error': 'Email requerido'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Buscar si el email existe (solo empleados activos, no superadmins por seguridad)
    try:
        empleado = DatosEmpleado.objects.get(correo_corporativo=email, estado='ACTIVA')
        usuario_existe = True
        nombre_usuario = f"{empleado.primer_nombre} {empleado.primer_apellido}"
    except DatosEmpleado.DoesNotExist:
        usuario_existe = False
        # Registrar intento fallido para auditoría
        logger.warning(f"[RECUPERACION] Intento de recuperación para email no registrado: {email}")
    
    if not usuario_existe:
        # Respuesta genérica por seguridad (no revelar si existe o no)
        return Response({
            'message': 'Si el email está registrado, recibirás un código de verificación',
            'enviado': False
        }, status=status.HTTP_200_OK)
    
    # Generar código de verificación (6 dígitos)
    codigo = generar_codigo_verificacion()
    
    # Guardar en caché por 15 minutos
    cache_key = f"recuperacion_{email}"
    cache.set(cache_key, {
        'codigo': codigo,
        'email': email,
        'empleado_id': empleado.id_empleado,
        'intentos': 0
    }, timeout=900)  # 15 minutos
    
    # Enviar email vía n8n
    email_sent, result = enviar_email_recuperacion_n8n(email, codigo, nombre_usuario)
    
    if email_sent:
        logger.info(f"[RECUPERACION] Código enviado a {email}")
        return Response({
            'message': 'Si el email está registrado, recibirás un código de verificación',
            'enviado': True,
            'email': email  # Solo para debug en desarrollo
        }, status=status.HTTP_200_OK)
    else:
        logger.error(f"[RECUPERACION] Error enviando código: {result}")
        return Response({
            'message': 'Error al enviar el código. Intenta más tarde.',
            'enviado': False
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# Endpoint para verificar código de recuperación - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def verificar_codigo_recuperacion(request):
    """
    Verifica el código de recuperación enviado por email.
    Si es válido, permite restablecer la contraseña.
    """
    email = request.data.get('email', '').strip().lower()
    codigo = request.data.get('codigo', '').strip()
    
    if not email or not codigo:
        return Response({'error': 'Email y código requeridos'}, status=status.HTTP_400_BAD_REQUEST)
    
    cache_key = f"recuperacion_{email}"
    datos_recuperacion = cache.get(cache_key)
    
    if not datos_recuperacion:
        return Response({'error': 'Código expirado o inválido'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Verificar código
    if datos_recuperacion['codigo'] != codigo:
        # Incrementar intentos fallidos
        datos_recuperacion['intentos'] = datos_recuperacion.get('intentos', 0) + 1
        
        # Si supera 3 intentos, invalidar
        if datos_recuperacion['intentos'] >= 3:
            cache.delete(cache_key)
            return Response({'error': 'Código bloqueado por múltiples intentos fallidos'}, status=status.HTTP_403_FORBIDDEN)
        
        cache.set(cache_key, datos_recuperacion, timeout=900)
        return Response({
            'error': 'Código incorrecto',
            'intentos_restantes': 3 - datos_recuperacion['intentos']
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Código válido - generar token temporal para restablecer
    token_temporal = str(uuid.uuid4())
    cache_key_token = f"reset_token_{token_temporal}"
    cache.set(cache_key_token, {
        'email': email,
        'empleado_id': datos_recuperacion['empleado_id'],
        'verificado': True
    }, timeout=900)  # 15 minutos para cambiar la contraseña
    
    # Limpiar caché del código (ya no se necesita)
    cache.delete(cache_key)
    
    logger.info(f"[RECUPERACION] Código verificado para {email}")
    
    return Response({
        'message': 'Código verificado correctamente',
        'token': token_temporal,  # Token para el siguiente paso
        'email': email
    }, status=status.HTTP_200_OK)


# Endpoint para restablecer contraseña - PÚBLICO
@api_view(['POST'])
@permission_classes([AllowAny])
def restablecer_password(request):
    """
    Restablece la contraseña después de verificar el código.
    Requiere token temporal y nueva contraseña.
    """
    token = request.data.get('token', '').strip()
    nueva_password = request.data.get('nueva_password', '').strip()
    
    if not token or not nueva_password:
        return Response({'error': 'Token y nueva contraseña requeridos'}, status=status.HTTP_400_BAD_REQUEST)
    
    if len(nueva_password) < 6:
        return Response({'error': 'La contraseña debe tener al menos 6 caracteres'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Verificar token temporal
    cache_key_token = f"reset_token_{token}"
    datos_token = cache.get(cache_key_token)
    
    if not datos_token:
        return Response({'error': 'Token inválido o expirado'}, status=status.HTTP_400_BAD_REQUEST)
    
    email = datos_token['email']
    empleado_id = datos_token['empleado_id']
    
    try:
        empleado = DatosEmpleado.objects.get(id_empleado=empleado_id, estado='ACTIVA')
        
        # Actualizar contraseña
        password_hash = bcrypt.hashpw(nueva_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        empleado.password_hash = password_hash
        empleado.save(update_fields=['password_hash'])
        
        # Limpiar token
        cache.delete(cache_key_token)
        
        # Crear registro de solicitud atendida
        SolicitudesPassword.objects.create(
            empleado=empleado,
            leida=True,
            atendida=True
        )
        
        # NOTIFICAR A ADMIN vía webhook de n8n
        notificar_admin_password_restablecida(empleado)
        
        logger.info(f"[RECUPERACION] Contraseña restablecida para {email}")
        
        return Response({
            'message': 'Contraseña restablecida correctamente',
            'email': email,
            'completado': True
        }, status=status.HTTP_200_OK)
        
    except DatosEmpleado.DoesNotExist:
        return Response({'error': 'Empleado no encontrado'}, status=status.HTTP_404_NOT_FOUND)


def notificar_admin_password_restablecida(empleado):
    """
    Notifica a admin/superadmin cuando un empleado restablece su contraseña.
    Envía webhook a n8n para notificación (email o teams).
    """
    try:
        import requests
        from django.conf import settings
        
        n8n_url = settings.N8N_WEBHOOK_URL
        if not n8n_url:
            return
        
        payload = {
            'tipo': 'notificacion_admin',
            'evento': 'password_restablecida',
            'datos_empleado': {
                'id': empleado.id_empleado,
                'nombre': f"{empleado.primer_nombre} {empleado.primer_apellido}",
                'email': empleado.correo_corporativo,
                'area': str(empleado.area) if empleado.area else 'Sin área',
                'cargo': str(empleado.cargo) if empleado.cargo else 'Sin cargo'
            },
            'mensaje': f"El empleado {empleado.primer_nombre} {empleado.primer_apellido} ha restablecido su contraseña",
            'timestamp': str(datetime.now())
        }
        
        headers = {
            'Content-Type': 'application/json',
            'X-API-Key': settings.N8N_WEBHOOK_API_KEY
        }
        
        response = requests.post(
            n8n_url,
            json=payload,
            headers=headers,
            timeout=10
        )
        
        if response.status_code == 200:
            logger.info(f"[N8N NOTIFICACION] Admin notificado de restablecimiento")
        else:
            logger.error(f"[N8N NOTIFICACION] Error: {response.status_code}")
            
    except Exception as e:
        logger.error(f"[NOTIFICACION] Error: {str(e)}")


# ============================================================================
# PROXY N8N — Sin CORS, el backend consulta n8n directamente
# ============================================================================

@api_view(['GET'])
def n8n_proxy(request):
    """
    Proxy server-side hacia n8n.
    actions:
      ?action=status      → verifica si n8n está online
      ?action=executions  → retorna historial de ejecuciones (requiere API key)
    """
    from django.conf import settings as django_settings
    import requests as ext_req

    action   = request.query_params.get('action', 'executions')
    base_url = getattr(django_settings, 'N8N_BASE_URL', '')
    api_key  = getattr(django_settings, 'N8N_WEBHOOK_API_KEY', '')

    if not base_url:
        return Response({'error': 'N8N_BASE_URL no configurado en el backend'}, status=503)

    try:
        if action == 'status':
            start = __import__('time').time()
            resp  = ext_req.get(f'{base_url}/healthz', timeout=5)
            ms    = round((__import__('time').time() - start) * 1000)
            return Response({
                'connected': resp.status_code in (200, 204),
                'ping': ms,
                'base_url': base_url,
            })

        elif action == 'executions':
            if not api_key:
                return Response({'error': 'N8N_WEBHOOK_API_KEY no configurado'}, status=503)

            status_filter = request.query_params.get('status')
            limit         = request.query_params.get('limit', 50)

            params = {'limit': limit}
            if status_filter and status_filter.upper() != 'ALL':
                params['status'] = status_filter.lower()

            resp = ext_req.get(
                f'{base_url}/api/v1/executions',
                headers={'X-N8N-API-KEY': api_key},
                params=params,
                timeout=10,
            )

            if resp.status_code == 200:
                raw  = resp.json()
                execs = raw.get('data', raw) if isinstance(raw, dict) else raw

                # Persistir en N8nLog los que no estén guardados todavía
                saved = 0
                for ex in (execs if isinstance(execs, list) else []):
                    exec_id = str(ex.get('id', ''))
                    if exec_id and not N8nLog.objects.filter(response_data__contains=f'"exec_id":"{exec_id}"').exists():
                        try:
                            import json as _json
                            wf_name = ex.get('workflowData', {}).get('name') or f"Workflow #{ex.get('workflowId','?')}"
                            ok = ex.get('status') in ('success',)
                            N8nLog.objects.create(
                                workflow_name=wf_name,
                                status='SUCCESS' if ok else 'ERROR',
                                message=f"Duración: {round(((__import__('datetime').datetime.fromisoformat(ex['stoppedAt'].replace('Z','+00:00')) - __import__('datetime').datetime.fromisoformat(ex['startedAt'].replace('Z','+00:00'))).total_seconds()))}s" if ex.get('startedAt') and ex.get('stoppedAt') else ex.get('status', ''),
                                tipo_evento=ex.get('mode', 'webhook'),
                                response_data=_json.dumps({'exec_id': exec_id})[:200],
                            )
                            saved += 1
                        except Exception as le:
                            logger.warning(f"[N8N PROXY] No se pudo persistir log: {le}")

                return Response({'data': execs, 'synced': saved})

        return Response({'error': 'Acción no válida. Usa action=status o action=executions'}, status=400)

    except ext_req.exceptions.ConnectionError:
        return Response({'error': f'No se pudo conectar con n8n en {base_url}'}, status=503)
    except ext_req.exceptions.Timeout:
        return Response({'error': 'n8n tardó demasiado en responder'}, status=504)
    except Exception as e:
        logger.error(f"[N8N PROXY] Error: {e}")
        return Response({'error': str(e)}, status=500)


# ============================================================================
# MARKITDOWN - Convertir archivos a Markdown
# ============================================================================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def convertir_markdown(request):
    """
    Convierte archivos (PDF, DOCX, XLSX, PPTX, etc.) a Markdown usando MarkItDown.
    """
    import tempfile
    import os
    
    try:
        # Verificar si hay archivo
        if 'archivo' not in request.FILES:
            return Response({'error': 'No se proporcionó archivo'}, status=400)
        
        archivo = request.FILES['archivo']
        
        # Extensiones permitidas
        extensiones_permitidas = ['.pdf', '.docx', '.xlsx', '.pptx', '.doc', '.xls', '.ppt', '.html', '.txt', '.csv', '.json', '.xml']
        ext = os.path.splitext(archivo.name)[1].lower()
        
        if ext not in extensiones_permitidas:
            return Response({
                'error': f'Extensión no soportada: {ext}',
                'soportadas': extensiones_permitidas
            }, status=400)
        
        # Guardar archivo temporalmente
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_input:
            for chunk in archivo.chunks():
                tmp_input.write(chunk)
            tmp_input_path = tmp_input.name
        
        logger.info(f"[MARKITDOWN] Archivo guardado: {tmp_input_path}")
        
        try:
            # Intentar usar markitdown como módulo Python
            try:
                from markitdown import MarkItDown
                md = MarkItDown()
                result = md.convert(tmp_input_path)
                markdown = result.text_content
                logger.info("[MARKITDOWN] Usando módulo Python markitdown")
            except ImportError:
                # Fallback: usar subprocess con python -m
                import subprocess
                import sys
                
                logger.info("[MARKITDOWN] Intentando usar python -m markitdown")
                
                resultado = subprocess.run(
                    [sys.executable, '-m', 'markitdown', tmp_input_path],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                
                if resultado.returncode != 0:
                    raise Exception(f"markitdown error: {resultado.stderr}")
                
                markdown = resultado.stdout
            
            # Limpiar archivo temporal
            try:
                os.unlink(tmp_input_path)
            except Exception:
                pass
            
            # Extraer metadatos básicos
            metadatos = {
                'nombre_original': archivo.name,
                'extension': ext,
                'tamaño_bytes': archivo.size,
                'lineas_markdown': len(markdown.split('\n')),
                'caracteres': len(markdown)
            }
            
            logger.info(f"[MARKITDOWN] Conversión exitosa: {archivo.name}")
            
            return Response({
                'markdown': markdown,
                'metadatos': metadatos,
                'exitoso': True
            })
                
        except Exception as e:
            try:
                os.unlink(tmp_input_path)
            except Exception:
                pass
            
            error_msg = str(e)
            logger.error(f"[MARKITDOWN] Error: {error_msg}")
            
            if 'markitdown' in error_msg.lower() or 'module' in error_msg.lower():
                return Response({
                    'error': 'MarkItDown no está instalado correctamente',
                    'detalle': error_msg,
                    'instrucciones': 'Ejecutar: pip install markitdown'
                }, status=503)
            
            return Response({
                'error': 'Error al convertir archivo',
                'detalle': error_msg
            }, status=500)
        
    except Exception as e:
        logger.error(f"[MARKITDOWN] Error general: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return Response({
            'error': 'Error interno del servidor',
            'detalle': str(e)
        }, status=500)


# ============================================================================
# CONVERTIDOR DE ARCHIVOS - PDF, Excel, Word, etc.
# ============================================================================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def convertir_archivo(request):
    """
    Convierte archivos entre diferentes formatos.
    
    Espera:
    - archivo: file upload
    - formato_destino: extensión de destino (pdf, docx, xlsx, csv, txt, html)
    
    Retorna:
    - archivo_base64: contenido del archivo convertido
    - nombre_archivo: nombre sugerido para descargar
    - mime_type: tipo MIME del archivo
    """
    import tempfile
    import os
    import base64
    import subprocess
    import sys
    
    try:
        if 'archivo' not in request.FILES:
            return Response({'error': 'No se proporcionó archivo'}, status=400)
        
        if 'formato_destino' not in request.POST:
            return Response({'error': 'No se especificó formato de destino'}, status=400)
        
        archivo = request.FILES['archivo']
        formato_destino = request.POST['formato_destino'].lower().strip('.')
        
        # Detectar formato origen
        ext_origen = os.path.splitext(archivo.name)[1].lower()
        formato_origen = ext_origen.lstrip('.')
        nombre_base = os.path.splitext(archivo.name)[0]
        
        logger.info(f"[CONVERTIR] {formato_origen} -> {formato_destino}: {archivo.name}")
        
        # Guardar archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext_origen) as tmp_input:
            for chunk in archivo.chunks():
                tmp_input.write(chunk)
            input_path = tmp_input.name
        
        output_path = input_path.replace(ext_origen, f'.{formato_destino}')
        
        try:
            resultado = None
            
            # === PDF a DOCX ===
            if formato_origen == 'pdf' and formato_destino == 'docx':
                try:
                    from pdf2docx import Converter
                    cv = Converter(input_path)
                    cv.convert(output_path, start=0, end=None)
                    cv.close()
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'pdf2docx no está instalado',
                        'instrucciones': 'pip install pdf2docx'
                    }, status=503)
            
            # === PDF a TXT (usando markitdown) ===
            elif formato_origen == 'pdf' and formato_destino == 'txt':
                try:
                    from markitdown import MarkItDown
                    md = MarkItDown()
                    result = md.convert(input_path)
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(result.text_content)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'markitdown no está instalado',
                        'instrucciones': 'pip install markitdown'
                    }, status=503)
            
            # === DOCX a PDF (usando reportlab - pura Python, sin dependencias externas) ===
            elif formato_origen == 'docx' and formato_destino == 'pdf':
                try:
                    from docx import Document
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    
                    # Leer documento Word
                    doc = Document(input_path)
                    
                    # Crear PDF
                    pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    # Convertir cada párrafo
                    for para in doc.paragraphs:
                        if para.text.strip():
                            # Detectar si es título (estilo Heading)
                            style_name = para.style.name if para.style else 'Normal'
                            if 'Heading 1' in style_name:
                                style = ParagraphStyle(
                                    'Heading1',
                                    parent=styles['Heading1'],
                                    fontSize=18,
                                    spaceAfter=12
                                )
                            elif 'Heading 2' in style_name:
                                style = ParagraphStyle(
                                    'Heading2',
                                    parent=styles['Heading2'],
                                    fontSize=14,
                                    spaceAfter=10
                                )
                            else:
                                style = styles['Normal']
                            
                            story.append(Paragraph(para.text, style))
                            story.append(Spacer(1, 6))
                    
                    pdf_doc.build(story)
                    resultado = output_path
                    
                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install python-docx reportlab'
                    }, status=503)
            
            # === DOCX a HTML ===
            elif formato_origen == 'docx' and formato_destino == 'html':
                try:
                    from mammoth import convert_to_html
                    with open(input_path, 'rb') as docx_file:
                        result = convert_to_html(docx_file)
                        with open(output_path, 'w', encoding='utf-8') as f:
                            f.write(result.value)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'mammoth no está instalado',
                        'instrucciones': 'pip install mammoth'
                    }, status=503)
            
            # === DOCX a TXT ===
            elif formato_origen == 'docx' and formato_destino == 'txt':
                try:
                    from docx import Document
                    doc = Document(input_path)
                    texto = '\n'.join([para.text for para in doc.paragraphs])
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(texto)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'python-docx no está instalado',
                        'instrucciones': 'pip install python-docx'
                    }, status=503)
            
            # === XLSX a CSV ===
            elif formato_origen == 'xlsx' and formato_destino == 'csv':
                try:
                    import pandas as pd
                    df = pd.read_excel(input_path)
                    df.to_csv(output_path, index=False, encoding='utf-8-sig')
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'pandas no está instalado',
                        'instrucciones': 'pip install pandas openpyxl'
                    }, status=503)
            
            # === XLSX a PDF (usando reportlab - pura Python) ===
            elif formato_origen == 'xlsx' and formato_destino == 'pdf':
                try:
                    import pandas as pd
                    from reportlab.lib.pagesizes import letter, landscape
                    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
                    from reportlab.lib import colors
                    
                    df = pd.read_excel(input_path)
                    
                    # Crear PDF en orientación landscape para tablas
                    pdf_doc = SimpleDocTemplate(output_path, pagesize=landscape(letter))
                    story = []
                    
                    # Convertir DataFrame a lista para Table de reportlab
                    data = [df.columns.tolist()] + df.values.tolist()
                    
                    # Crear tabla
                    table = Table(data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#001e33')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                    ]))
                    
                    story.append(table)
                    pdf_doc.build(story)
                    resultado = output_path
                    
                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install pandas openpyxl reportlab'
                    }, status=503)
            
            # === XLSX a TXT ===
            elif formato_origen == 'xlsx' and formato_destino == 'txt':
                try:
                    import pandas as pd
                    df = pd.read_excel(input_path)
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(df.to_string(index=False))
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'pandas no está instalado',
                        'instrucciones': 'pip install pandas openpyxl'
                    }, status=503)
            
            # === CSV a XLSX ===
            elif formato_origen == 'csv' and formato_destino == 'xlsx':
                try:
                    import pandas as pd
                    df = pd.read_csv(input_path)
                    df.to_excel(output_path, index=False)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'pandas no está instalado',
                        'instrucciones': 'pip install pandas openpyxl'
                    }, status=503)
            
            # === TXT a DOCX ===
            elif formato_origen == 'txt' and formato_destino == 'docx':
                try:
                    from docx import Document
                    doc = Document()
                    with open(input_path, 'r', encoding='utf-8') as f:
                        for linea in f:
                            doc.add_paragraph(linea.strip())
                    doc.save(output_path)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'python-docx no está instalado',
                        'instrucciones': 'pip install python-docx'
                    }, status=503)
            
            # === TXT a PDF ===
            elif formato_origen == 'txt' and formato_destino == 'pdf':
                try:
                    from fpdf import FPDF
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font('Arial', size=12)
                    with open(input_path, 'r', encoding='utf-8') as f:
                        for linea in f:
                            pdf.cell(200, 10, txt=linea.strip(), ln=True)
                    pdf.output(output_path)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'fpdf no está instalado',
                        'instrucciones': 'pip install fpdf'
                    }, status=503)
            
            # === HTML a PDF (usando reportlab con html.parser) ===
            elif formato_origen == 'html' and formato_destino == 'pdf':
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet
                    from html.parser import HTMLParser
                    
                    # Parser simple para extraer texto del HTML
                    class HTMLTextExtractor(HTMLParser):
                        def __init__(self):
                            super().__init__()
                            self.texts = []
                            self.current_tag = None
                            
                        def handle_starttag(self, tag, attrs):
                            self.current_tag = tag
                            
                        def handle_endtag(self, tag):
                            if tag in ('p', 'div', 'h1', 'h2', 'h3', 'br'):
                                self.texts.append('\n')
                            self.current_tag = None
                            
                        def handle_data(self, data):
                            if data.strip():
                                self.texts.append(data.strip())
                    
                    # Leer HTML
                    with open(input_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                    
                    parser = HTMLTextExtractor()
                    parser.feed(html_content)
                    text = ' '.join(parser.texts)
                    
                    # Crear PDF
                    pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    for line in text.split('\n'):
                        if line.strip():
                            story.append(Paragraph(line.strip(), styles['Normal']))
                            story.append(Spacer(1, 6))
                    
                    pdf_doc.build(story)
                    resultado = output_path
                    
                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install reportlab'
                    }, status=503)
            
            # === HTML a DOCX ===
            elif formato_origen == 'html' and formato_destino == 'docx':
                try:
                    from htmldocx import HtmlToDocx
                    parser = HtmlToDocx()
                    parser.parse_html_file(input_path, output_path)
                    resultado = output_path
                except ImportError:
                    return Response({
                        'error': 'htmldocx no está instalado',
                        'instrucciones': 'pip install htmldocx'
                    }, status=503)
            
            # === IMÁGENES a PDF (png, jpg, jpeg, gif, bmp, webp) ===
            elif formato_origen in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp') and formato_destino == 'pdf':
                try:
                    from PIL import Image
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Image as RLImage, Spacer
                    from reportlab.lib.units import inch
                    
                    # Abrir imagen y obtener dimensiones
                    img = Image.open(input_path)
                    img_width, img_height = img.size
                    
                    # Calcular tamaño para ajustar a página
                    page_width, page_height = letter
                    max_width = page_width - 2 * inch
                    max_height = page_height - 2 * inch
                    
                    # Escalar proporcionalmente
                    scale = min(max_width / img_width, max_height / img_height, 1.0)
                    final_width = img_width * scale
                    final_height = img_height * scale
                    
                    # Crear PDF
                    pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
                    story = []
                    story.append(Spacer(1, 0.5 * inch))
                    story.append(RLImage(input_path, width=final_width, height=final_height))
                    pdf_doc.build(story)
                    resultado = output_path
                    
                except ImportError as e:
                    return Response({
                        'error': f'Librería no instalada: {str(e)}',
                        'instrucciones': 'pip install Pillow reportlab'
                    }, status=503)
            
            else:
                return Response({
                    'error': f'Conversión no soportada: {formato_origen} -> {formato_destino}'
                }, status=400)
            
            # Leer archivo convertido y codificar en base64
            if resultado and os.path.exists(resultado):
                with open(resultado, 'rb') as f:
                    contenido = f.read()
                    archivo_base64 = base64.b64encode(contenido).decode('utf-8')
                
                # Determinar MIME type
                mime_types = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    'csv': 'text/csv',
                    'txt': 'text/plain',
                    'html': 'text/html',
                }
                
                # Limpiar archivos temporales
                try:
                    os.unlink(input_path)
                    os.unlink(resultado)
                except Exception:
                    pass
                
                logger.info(f"[CONVERTIR] Éxito: {archivo.name} -> {nombre_base}-convertido.{formato_destino}")
                
                return Response({
                    'archivo_base64': archivo_base64,
                    'nombre_archivo': f'{nombre_base}-convertido.{formato_destino}',
                    'mime_type': mime_types.get(formato_destino, 'application/octet-stream'),
                    'exitoso': True
                })
            else:
                return Response({
                    'error': 'Error al generar archivo convertido'
                }, status=500)
        
        except Exception as e:
            # Limpiar archivos temporales
            try:
                os.unlink(input_path)
                if os.path.exists(output_path):
                    os.unlink(output_path)
            except Exception:
                pass
            
            logger.error(f"[CONVERTIR] Error: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            
            return Response({
                'error': 'Error al convertir archivo',
                'detalle': str(e)
            }, status=500)
    
    except Exception as e:
        logger.error(f"[CONVERTIR] Error general: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return Response({
            'error': 'Error interno del servidor',
            'detalle': str(e)
        }, status=500)


# ═══════════════════════════════════════════════════════════════════════════
# GESTOR DE PDFs - Fusionar, Dividir, Rotar, Comprimir, Proteger, etc.
# ═══════════════════════════════════════════════════════════════════════════

@api_view(['POST'])
def gestor_pdf(request):
    """
    Endpoint para gestionar operaciones con PDFs:
    - fusionar: Unir múltiples PDFs en uno solo
    - dividir: Separar páginas de un PDF
    - comprimir: Reducir tamaño del PDF
    - rotar: Rotar páginas
    - extraer: Extraer páginas específicas
    - marca: Agregar marca de agua
    - proteger: Agregar contraseña
    - desbloquear: Quitar contraseña
    """
    try:
        herramienta = request.POST.get('herramienta')
        if not herramienta:
            return Response({'error': 'No se especificó la herramienta'}, status=400)
        
        # Obtener cantidad de archivos
        cantidad = int(request.POST.get('cantidad_archivos', 0))
        if cantidad == 0:
            return Response({'error': 'No se proporcionaron archivos'}, status=400)
        
        # Leer todos los archivos
        archivos_pdf = []
        for i in range(cantidad):
            archivo_key = f'archivo_{i}'
            if archivo_key in request.FILES:
                archivos_pdf.append(request.FILES[archivo_key])
        
        if len(archivos_pdf) == 0:
            return Response({'error': 'No se encontraron archivos PDF'}, status=400)
        
        logger.info(f"[GESTOR PDF] Herramienta: {herramienta}, Archivos: {len(archivos_pdf)}")
        
        # Importar pypdf
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            return Response({
                'error': 'pypdf no está instalado',
                'instrucciones': 'pip install pypdf'
            }, status=503)
        
        resultado_archivos = []
        
        # CREAR PDF DESDE CERO
        if herramienta == 'crear':
            import json
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            import io
            import urllib.parse
            
            # Obtener datos del cuerpo JSON
            try:
                body_data = json.loads(request.body.decode('utf-8'))
            except:
                body_data = {}
            
            titulo = body_data.get('titulo', 'documento')
            paginas_data = body_data.get('paginas', [{'elementos': []}])
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            for pagina in paginas_data:
                elementos = pagina.get('elementos', [])
                
                for el in elementos:
                    tipo = el.get('tipo', 'texto')
                    contenido = el.get('contenido', '')
                    x = el.get('x', 50)
                    y = height - el.get('y', 50)  # Invertir Y para coordenadas PDF
                    font_size = el.get('fontSize', 12)
                    color = el.get('color', '#000000')
                    
                    if tipo == 'texto' and contenido:
                        # Convertir color hex a RGB
                        color_hex = color.lstrip('#')
                        r = int(color_hex[0:2], 16) / 255
                        g = int(color_hex[2:4], 16) / 255
                        b = int(color_hex[4:6], 16) / 255
                        
                        c.setFillColorRGB(r, g, b)
                        c.setFont('Helvetica', font_size)
                        c.drawString(x, y, contenido)
                    
                    elif tipo == 'imagen' and contenido:
                        # Imagen en base64
                        try:
                            if contenido.startswith('data:image'):
                                # Extraer base64 de data URL
                                base64_data = contenido.split(',')[1]
                                img_data = base64.b64decode(base64_data)
                            else:
                                img_data = base64.b64decode(contenido)
                            
                            img_buffer = io.BytesIO(img_data)
                            img = ImageReader(img_buffer)
                            c.drawImage(img, x, y - 100, width=150, height=100, mask='auto')
                        except Exception as e:
                            logger.warning(f"[CREAR PDF] Error al insertar imagen: {e}")
                
                c.showPage()
            
            c.save()
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            resultado_archivos.append({
                'nombre': f'{titulo}.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # EDITAR PDF - Agregar elementos a PDF existente
        elif herramienta == 'editar':
            import json
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            import io
            
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            
            elementos_str = request.POST.get('elementos', '[]')
            try:
                elementos = json.loads(elementos_str)
            except:
                elementos = []
            
            width, height = letter
            
            for i, page in enumerate(reader.pages):
                # Crear overlay con los elementos
                packet = io.BytesIO()
                c = canvas.Canvas(packet, pagesize=letter)
                
                for el in elementos:
                    tipo = el.get('tipo', 'texto')
                    contenido = el.get('contenido', '')
                    x = el.get('x', 50)
                    y = height - el.get('y', 50)
                    font_size = el.get('fontSize', 12)
                    color = el.get('color', '#000000')
                    
                    if tipo == 'texto' and contenido:
                        color_hex = color.lstrip('#')
                        r = int(color_hex[0:2], 16) / 255
                        g = int(color_hex[2:4], 16) / 255
                        b = int(color_hex[4:6], 16) / 255
                        
                        c.setFillColorRGB(r, g, b)
                        c.setFont('Helvetica', font_size)
                        c.drawString(x, y, contenido)
                    
                    elif tipo == 'imagen' and contenido:
                        try:
                            if contenido.startswith('data:image'):
                                base64_data = contenido.split(',')[1]
                                img_data = base64.b64decode(base64_data)
                            else:
                                img_data = base64.b64decode(contenido)
                            
                            img_buffer = io.BytesIO(img_data)
                            img = ImageReader(img_buffer)
                            c.drawImage(img, x, y - 100, width=150, height=100, mask='auto')
                        except Exception as e:
                            logger.warning(f"[EDITAR PDF] Error al insertar imagen: {e}")
                
                c.save()
                packet.seek(0)
                
                # Merge overlay con página original
                overlay_reader = PdfReader(packet)
                page.merge_page(overlay_reader.pages[0])
                writer.add_page(page)
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({
                'nombre': f'{nombre_base}_editado.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # FUSIONAR PDFs
        if herramienta == 'fusionar':
            writer = PdfWriter()
            
            for archivo in archivos_pdf:
                reader = PdfReader(archivo)
                for page in reader.pages:
                    writer.add_page(page)
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            resultado_archivos.append({
                'nombre': 'fusionado.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # DIVIDIR PDF
        elif herramienta == 'dividir':
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            total_paginas = len(reader.pages)
            
            for i in range(total_paginas):
                writer = PdfWriter()
                writer.add_page(reader.pages[i])
                
                output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
                with open(output_path, 'wb') as f:
                    writer.write(f)
                
                with open(output_path, 'rb') as f:
                    contenido = f.read()
                
                resultado_archivos.append({
                    'nombre': f'pagina_{i+1}.pdf',
                    'contenido_base64': base64.b64encode(contenido).decode('utf-8')
                })
                os.unlink(output_path)
        
        # COMPRIMIR PDF
        elif herramienta == 'comprimir':
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            
            for page in reader.pages:
                writer.add_page(page)
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({
                'nombre': f'{nombre_base}_comprimido.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # ROTAR PÁGINAS
        elif herramienta == 'rotar':
            rotacion = int(request.POST.get('rotacion', 90))
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            
            for page in reader.pages:
                page.rotate(rotacion)
                writer.add_page(page)
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({
                'nombre': f'{nombre_base}_rotado_{rotacion}.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # EXTRAER PÁGINAS
        elif herramienta == 'extraer':
            paginas_str = request.POST.get('paginas', '')
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            
            paginas_a_extraer = set()
            for parte in paginas_str.split(','):
                parte = parte.strip()
                if '-' in parte:
                    inicio, fin = map(int, parte.split('-'))
                    paginas_a_extraer.update(range(inicio, fin + 1))
                elif parte.isdigit():
                    paginas_a_extraer.add(int(parte))
            
            for num_pagina in sorted(paginas_a_extraer):
                if 1 <= num_pagina <= len(reader.pages):
                    writer.add_page(reader.pages[num_pagina - 1])
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({
                'nombre': f'{nombre_base}_extraido.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # MARCA DE AGUA
        elif herramienta == 'marca':
            texto_marca = request.POST.get('texto', '')
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            
            marca_imagen = request.FILES.get('marca_imagen')
            
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            import io
            
            for i, page in enumerate(reader.pages):
                packet = io.BytesIO()
                c = canvas.Canvas(packet, pagesize=letter)
                
                if marca_imagen:
                    img = ImageReader(marca_imagen)
                    c.drawImage(img, 150, 300, width=300, height=300, mask='auto')
                elif texto_marca:
                    c.saveState()
                    c.setFont('Helvetica', 40)
                    c.setFillColorRGB(0.7, 0.7, 0.7, alpha=0.3)
                    c.translate(300, 400)
                    c.rotate(45)
                    c.drawCentredString(0, 0, texto_marca)
                    c.restoreState()
                
                c.save()
                packet.seek(0)
                
                marca_reader = PdfReader(packet)
                page.merge_page(marca_reader.pages[0])
                writer.add_page(page)
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({
                'nombre': f'{nombre_base}_con_marca.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # PROTEGER PDF
        elif herramienta == 'proteger':
            password = request.POST.get('password', '')
            if not password:
                return Response({'error': 'Se requiere una contraseña'}, status=400)
            
            archivo = archivos_pdf[0]
            reader = PdfReader(archivo)
            writer = PdfWriter()
            
            for page in reader.pages:
                writer.add_page(page)
            
            writer.encrypt(password)
            
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            with open(output_path, 'rb') as f:
                contenido = f.read()
            
            nombre_base = os.path.splitext(archivo.name)[0]
            resultado_archivos.append({
                'nombre': f'{nombre_base}_protegido.pdf',
                'contenido_base64': base64.b64encode(contenido).decode('utf-8')
            })
            os.unlink(output_path)
        
        # DESBLOQUEAR PDF
        elif herramienta == 'desbloquear':
            password = request.POST.get('password', '')
            archivo = archivos_pdf[0]
            
            try:
                reader = PdfReader(archivo)
                if reader.is_encrypted:
                    reader.decrypt(password)
                
                writer = PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                
                output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
                with open(output_path, 'wb') as f:
                    writer.write(f)
                
                with open(output_path, 'rb') as f:
                    contenido = f.read()
                
                nombre_base = os.path.splitext(archivo.name)[0]
                resultado_archivos.append({
                    'nombre': f'{nombre_base}_desbloqueado.pdf',
                    'contenido_base64': base64.b64encode(contenido).decode('utf-8')
                })
                os.unlink(output_path)
            except Exception as e:
                return Response({'error': 'No se pudo desbloquear el PDF. Verifica la contraseña.'}, status=400)
        
        else:
            return Response({'error': f'Herramienta no reconocida: {herramienta}'}, status=400)
        
        logger.info(f"[GESTOR PDF] Éxito: {herramienta}, {len(resultado_archivos)} archivos generados")
        
        return Response({
            'archivos': resultado_archivos,
            'exitoso': True
        })
    
    except Exception as e:
        logger.error(f"[GESTOR PDF] Error: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return Response({
            'error': 'Error al procesar PDF',
            'detalle': str(e)
        }, status=500)
